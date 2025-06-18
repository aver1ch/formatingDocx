from docxtpl import DocxTemplate, InlineImage
from docxcompose.composer import Composer
from docx.shared import Mm
from docx import Document

def add_title_page_to_document(template_path, image_path, source_doc_path, output_path):
    # Рендерим титульный лист
    title_doc = DocxTemplate(template_path)
    context = {
        'agency_name': 'ФЕДЕРАЛЬНОЕ АГЕНТСТВО\nПО ТЕХНИЧЕСКОМУ РЕГУЛИРОВАНИЮ И МЕТРОЛОГИИ',
        'standart_type': 'ХУЕСТ Р',
        'image': InlineImage(title_doc, image_path, width=Mm(42)),
        'designation': 'че-то про дизайн',
        'title': 'Доклад о Python. Про моего питона',
        'status': 'Редакционный проект',
        'city': 'Москва',
        'publisher_info': 'Российский институт стандартизации',
        'current_year': '2025',
    }
    title_doc.render(context)
    title_doc.save("temp_title.docx")

    # Объединяем документы
    composer = Composer(Document())
    composer.append(Document("temp_title.docx"))
    composer.append(Document(source_doc_path))
    composer.save(output_path)

# Пример использования
add_title_page_to_document(
    template_path="doc_editor/templates/title_page_template.docx",
    image_path="doc_editor/doc_builder/logo.png",
    source_doc_path="doc_editor/tests/test_data/test.docx",
    output_path="final2.docx"
)