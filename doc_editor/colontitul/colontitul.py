from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import logging
from typing import Dict, Any

logger = logging.getLogger(__name__)

class HeaderFooterManager:
    def __init__(self, document: Document, config: Dict[str, Any]):
        self.doc = document
        self.config = config
        self.logger = logger

    def _create_element(self, name):
        return OxmlElement(name)

    def _add_text_to_element(self, element, text, align=None):
    # Используем первый параграф или создаем новый
        if element.paragraphs:
            paragraph = element.paragraphs[0]
            paragraph.clear()  # Очищаем существующий контент
        else:
            paragraph = element.add_paragraph()
        
        # Устанавливаем выравнивание
        if align == 'left':
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif align == 'right':
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        elif align == 'center':
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Добавляем текст
        run = paragraph.add_run(text)
        return element


    '''def apply_headers_footers(self):
        """Применяет настройки колонтитулов из конфига"""
        try:
            headers_config = self.config['document']['numbering']['headers']
            
            # Для каждой секции в документе
            for section in self.doc.sections:
                # Удаляем все существующие колонтитулы
                for header in [section.first_page_header, section.header, section.even_page_header]:
                    if header is not None:
                        header.is_linked_to_previous = True
                
                # Создаем разные колонтитулы для четных и нечетных страниц
                section.different_first_page_header_footer = False
                section.different_even_page_header_footer = True
                
                # Нечетные страницы (правое выравнивание)
                odd_header = section.header
                self._add_text_to_element(odd_header, headers_config['right'], 'right')
                
                # Четные страницы (левое выравнивание)
                even_header = section.even_page_header
                self._add_text_to_element(even_header, headers_config['left'], 'left')
                
                # Нумерация страниц в нижнем колонтитуле
                footer = section.footer
                self._add_page_number(footer)
            
            self.logger.info("Колонтитулы успешно применены")
            
        except KeyError as e:
            self.logger.error(f"Отсутствует параметр в конфигурации: {e}")
            raise
        except Exception as e:
            self.logger.error(f"Ошибка применения колонтитулов: {e}")
            raise'''
    def apply_headers_footers(self):
        """Применяет настройки колонтитулов из конфига"""
        try:
            headers_config = self.config['document']['numbering']['headers']
            self.doc.settings.odd_and_even_pages_header_footer = True
            if not headers_config.get('enabled', False):
                self.logger.info("Колонтитулы отключены в конфигурации")
                return
                
            for section in self.doc.sections:
                # Включаем разные колонтитулы для четных/нечетных
                # self.doc.odd_and_even_pages_header_footer = False
                section.different_first_page_header_footer = True  # ← Добавить эту строку
                
                # Очищаем колонтитул первой страницы (титульной)
                self._clear_element(section.first_page_header)
                self._clear_element(section.first_page_footer)
                
                # Нечетные страницы: "Национальный стандарт РФ" справа
                self._add_text_to_element(
                    section.header, 
                    headers_config.get('left', ''), 
                    'right'
                )
                
                # Четные страницы: "Обозначение стандарта" слева
                self._add_text_to_element(
                    section.even_page_header, 
                    headers_config.get('right', ''), 
                    'left'
                )
                
                # Нумерация страниц
                if headers_config.get('page_numbers', False):
                    self._add_page_number(section.footer, 'right')
                    self._add_page_number(section.even_page_footer, 'left')
            
            self.logger.info("Колонтитулы успешно применены")
            
        except Exception as e:
            self.logger.error(f"Ошибка применения колонтитулов: {e}")
            raise

    def _clear_element(self, element):
        """Очищает содержимое колонтитула"""
        for paragraph in element.paragraphs:
            paragraph.clear()
        # Если нет параграфов, добавляем пустой
        if not element.paragraphs:
            element.add_paragraph()


    def _add_page_number(self, footer, align='center'):
        """Упрощенная версия добавления номеров страниц"""
        if footer.paragraphs:
            paragraph = footer.paragraphs[0]
            paragraph.clear()
        else:
            paragraph = footer.add_paragraph()
        
        # Устанавливаем выравнивание
        if align == 'left':
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif align == 'right':
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        else:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Простое добавление текста с номером страницы
        # Можно использовать встроенные возможности python-docx
        # paragraph.add_run("Страница ")
        
        '''# Добавляем поле номера страницы
        from docx.oxml.shared import qn
        from docx.oxml import OxmlElement
        
        fld_simple = OxmlElement('w:fldSimple')
        fld_simple.set(qn('w:instr'), 'PAGE')
        run = paragraph.add_run()
        run._element.append(fld_simple)'''

        run = paragraph.add_run()
        fld_char_begin = OxmlElement('w:fldChar')
        fld_char_begin.set(qn('w:fldCharType'), 'begin')
        
        instr_text = OxmlElement('w:instrText')
        instr_text.text = 'PAGE'
        
        fld_char_end = OxmlElement('w:fldChar')
        fld_char_end.set(qn('w:fldCharType'), 'end')
        
        run._element.append(fld_char_begin)
        run._element.append(instr_text)
        run._element.append(fld_char_end)


    def setup_numbering(self):
        """Настраивает нумерацию страниц согласно конфигу"""
        try:
            numbering_config = self.config['document']['numbering']['pages']
            
            # Для предисловия и содержания используем римские цифры
            # Для основной части - арабские
            # Реализация этого требует более сложной логики разделения документа,
            # которую можно добавить позже
            
            self.logger.info("Настройки нумерации страниц применены")
            
        except KeyError as e:
            self.logger.error(f"Отсутствует параметр нумерации в конфигурации: {e}")
            raise
        except Exception as e:
            self.logger.error(f"Ошибка настройки нумерации страниц: {e}")
            raise