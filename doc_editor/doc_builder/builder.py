from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Dict, Any
import re
from datetime import datetime
import logging
from doc_editor.doc_builder.template_manager import TemplateManager

logger = logging.getLogger(__name__)

class DocumentStructureBuilder:
    def __init__(self, doc: Document, config: Dict[str, Any]):
        """
        Инициализация построителя структуры документа.

        :param doc: Объект документа docx
        :param config: Конфигурация из YAML (полный словарь)
        """
        self.doc = doc
        self.config = config
        self._current_section = self.doc.sections[0]
        self.template_manager = TemplateManager()
        self.logger = logging.getLogger(f"{__name__}.{self.__class__.__name__}")

        self._validate_structure_config()
        self._setup_title_styles()

    def _validate_structure_config(self) -> None:
        """Проверка наличия обязательных полей в конфиге структуры"""
        if 'document' not in self.config:
            raise ValueError("Missing 'document' section in config")

        if 'structure' not in self.config['document']:
            raise ValueError("Missing 'structure' section in document config")

        if 'title_page' not in self.config['document']['structure']:
            self.config['document']['structure']['title_page'] = {'elements': []}

    def _setup_title_styles(self) -> None:
        """Создаем кастомные стили для титульной страницы"""
        styles = {
            'Custom_Title_Agency': {'size': 12, 'bold': True, 'alignment': WD_ALIGN_PARAGRAPH.CENTER},
            'Custom_Title_StandardType': {'size': 14, 'bold': True, 'alignment': WD_ALIGN_PARAGRAPH.CENTER},
            'Custom_Title_Designation': {'size': 16, 'bold': True, 'alignment': WD_ALIGN_PARAGRAPH.CENTER,
                                         'spacing_after': Pt(24)},
            'Custom_Title_Main': {'size': 18, 'bold': True, 'alignment': WD_ALIGN_PARAGRAPH.CENTER,
                                  'spacing_before': Pt(36), 'spacing_after': Pt(36)},
            'Custom_Title_Status': {'size': 12, 'italic': True, 'alignment': WD_ALIGN_PARAGRAPH.CENTER,
                                    'spacing_before': Pt(24)},
            'Custom_Title_Publisher': {'size': 11, 'alignment': WD_ALIGN_PARAGRAPH.CENTER, 'spacing_before': Pt(48)}
        }

        for style_name, style_props in styles.items():
            if style_name not in self.doc.styles:
                style = self.doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
                font = style.font
                font.name = self.config['document']['general']['fonts']['main']['family']
                font.size = Pt(style_props['size'])
                if 'bold' in style_props:
                    font.bold = style_props['bold']
                if 'italic' in style_props:
                    font.italic = style_props['italic']

                paragraph_format = style.paragraph_format
                paragraph_format.alignment = style_props['alignment']
                if 'spacing_before' in style_props:
                    paragraph_format.space_before = style_props['spacing_before']
                if 'spacing_after' in style_props:
                    paragraph_format.space_after = style_props['spacing_after']

    def _clear_header_footer(self) -> None:
        """Очистка колонтитулов для текущей секции"""

        def clear_element(element):
            if element is not None:
                for paragraph in element.paragraphs:
                    p = paragraph._element
                    p.getparent().remove(p)
                for table in element.tables:
                    tbl = table._element
                    tbl.getparent().remove(tbl)

        clear_element(self._current_section.header)
        clear_element(self._current_section.first_page_header)
        clear_element(self._current_section.footer)
        clear_element(self._current_section.first_page_footer)

    def _get_title_page_element(self, element_name: str) -> Any:
        """Получает значение элемента титульной страницы из конфига"""
        for element in self.config['document']['structure']['title_page']['elements']:
            if element_name in element:
                return element[element_name]
        return None

    def _configure_title_page_section(self) -> None:
        """Настройка параметров секции титульной страницы"""
        # Используйте поля из конфигурации
        margins_config = self.config['document']['general']['margins']
        self._current_section.top_margin = self._parse_measurement(margins_config['top'])
        self._clear_header_footer()

    def _parse_measurement(self, value: str):
        """Парсинг размеров (копия из editor.py)"""
        from docx.shared import Mm, Cm, Pt
        
        if isinstance(value, (int, float)):
            return Pt(float(value))
        
        value = str(value).strip().lower()
        
        if value.endswith('mm'):
            return Mm(float(value[:-2]))
        elif value.endswith('cm'):
            return Cm(float(value[:-2]))
        elif value.endswith('pt'):
            return Pt(float(value[:-2]))
        else:
            return Pt(float(value))

    def _add_agency_name(self, text: str) -> None:
        """Добавление названия агентства"""
        p = self.doc.add_paragraph(text, style='Custom_Title_Agency')
        self.doc.add_paragraph()  # Пустой абзац для отступа

    def _add_logo_placeholder(self, position: str) -> None:
        """Добавление placeholder для логотипа"""
        align_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT
        }

        p = self.doc.add_paragraph()
        p.alignment = align_map.get(position.lower(), WD_ALIGN_PARAGRAPH.CENTER)
        run = p.add_run()
        # В реальном коде замените на путь к вашему логотипу
        run.add_picture('doc_editor/doc_builder/logo.png', width=Cm(2))
        self.doc.add_paragraph()  # Отступ после лого

    def _add_standard_info(self, standard_type: str, designation: str) -> None:
        """Добавление информации о стандарте"""
        p_type = self.doc.add_paragraph(standard_type, style='Custom_Title_StandardType')
        p_desig = self.doc.add_paragraph(designation, style='Custom_Title_Designation')

    def _add_document_title(self, title: str) -> None:
        """Добавление названия документа"""
        p = self.doc.add_paragraph(title, style='Custom_Title_Main')
        self._add_horizontal_line()

    def _add_horizontal_line(self) -> None:
        """Добавление горизонтальной линии под заголовком"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Простой способ добавить горизонтальную линию
        run = p.add_run()
        run.add_break()
        run.add_break()
        run.add_break(WD_BREAK.LINE)
        run.add_break()
        run.add_break()

    def _add_status_info(self, status: str) -> None:
        """Добавление информации о статусе документа"""
        self.doc.add_paragraph(status, style='Custom_Title_Status')

    def _add_publisher_info(self, info_template: str) -> None:
        """Добавление информации об издателе"""
        info = info_template.replace('год', str(datetime.now().year))
        self.doc.add_paragraph(info, style='Custom_Title_Publisher')

    def _add_section_break(self) -> None:
        """Добавление разрыва раздела"""
        self.doc.add_paragraph()
        self.doc.add_page_break()
        self._current_section = self.doc.add_section()

    def build_document_structure(self) -> None:
        """Построение всей структуры документа"""
        try:
            # Вставляем титульник В НАЧАЛО документа из шаблона
            if self._should_add_title_page():
                success = self._insert_title_page_from_template()
                if not success:
                    self.logger.warning("Не удалось вставить титульник из шаблона, используем fallback")
                    self._build_title_page_fallback()
            
            # Добавляем предисловие (если нужно)
            if self._should_add_preface():
                self._build_preface()
                self._add_section_break()
                
            self.logger.info("Структура документа успешно построена")
            
        except Exception as e:
            self.logger.error(f"Ошибка построения структуры: {e}")
            raise

    def _build_title_page_fallback(self) -> None:
        """Резервный метод создания титульника (старая логика)"""
        self.logger.info("Используется резервный метод создания титульника")
        self._build_title_page()


    def _insert_title_page_from_template(self) -> bool:
        """
        Вставка титульного листа из шаблона
        
        :return: True если успешно, False если нужен fallback
        """
        try:
            # Получаем имя шаблона из конфигурации
            template_name = self.config['document']['structure']['title_page'].get(
                'template', 'title_page_template'
            )
            
            # Загружаем шаблон
            template_doc = self.template_manager.load_template(template_name)
            if not template_doc:
                return False
            
            # Подготавливаем данные для замены плейсхолдеров
            placeholder_data = self._prepare_title_page_data()
            
            # Заменяем плейсхолдеры
            template_doc = self.template_manager.replace_placeholders(template_doc, placeholder_data)
            
            # Вставляем в начало документа
            self.template_manager.insert_template_content(self.doc, template_doc)
            
            self.logger.info(f"Титульный лист успешно вставлен из шаблона: {template_name}")
            return True
            
        except Exception as e:
            self.logger.error(f"Ошибка вставки титульного листа из шаблона: {e}")
            return False
        
    def _prepare_title_page_data(self) -> Dict[str, str]:
        """Подготовка данных для замены плейсхолдеров в титульном листе"""
        data = {}
        
        # Получаем элементы титульного листа из конфигурации
        title_page_elements = self.config['document']['structure']['title_page'].get('elements', [])
        
        # Преобразуем список словарей в единый словарь
        for element in title_page_elements:
            if isinstance(element, dict):
                for key, value in element.items():
                    data[key] = str(value) if value is not None else ""
        
        # Добавляем текущий год
        data['current_year'] = str(datetime.now().year)
        
        # Устанавливаем значения по умолчанию для пустых полей
        default_values = {
            'agency_name': '',
            'standard_type': '',
            'designation': '',
            'title': '',
            'status': '',
            'publisher_info': ''
        }
        
        for key, default_value in default_values.items():
            if key not in data:
                data[key] = default_value
        
        self.logger.info(f"Подготовлены данные для замены: {list(data.keys())}")
        return data

    def _should_add_title_page(self) -> bool:
        """Нужно ли добавлять титульную страницу?"""
        title_page_config = self.config['document']['structure'].get('title_page', {})
        # ✅ Проверяем наличие шаблона ИЛИ элементов
        return bool(title_page_config.get('elements', []) or title_page_config.get('template'))

    def _should_add_preface(self) -> bool:
        """Нужно ли добавлять предисловие?"""
        preface_cfg = self.config['document']['structure'].get('preface', {})
        return preface_cfg.get('required', False)

    def _build_title_page(self) -> None:
        """Построение титульной страницы"""
        self._configure_title_page_section()

        agency_name = self._get_title_page_element('agency_name')
        if agency_name:
            self._add_agency_name(agency_name)

        logo_pos = self._get_title_page_element('logo_position')
        if logo_pos:
            self._add_logo_placeholder(logo_pos)

        standard_type = self._get_title_page_element('standard_type')
        designation = self._get_title_page_element('designation')
        if standard_type and designation:
            self._add_standard_info(standard_type, designation)

        title = self._get_title_page_element('title')
        if title:
            self._add_document_title(title)

        status = self._get_title_page_element('status')
        if status:
            self._add_status_info(status)

        publisher_info = self._get_title_page_element('publisher_info')
        if publisher_info:
            self._add_publisher_info(publisher_info)

    def _build_preface(self) -> None:
        """Построение предисловия"""
        self._add_preface_title()

        for content_type in self.config['document']['structure']['preface'].get('content', []):
            if content_type == 'development_info':
                self._add_development_info()
            elif content_type == 'approval_info':
                self._add_approval_info()
            elif content_type == 'replacement_info':
                self._add_replacement_info()
            elif content_type == 'patent_notice':
                self._add_patent_notice()

    def _add_preface_title(self) -> None:
        """Добавление заголовка предисловия"""
        p = self.doc.add_paragraph('Предисловие', style='Heading 1')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_development_info(self) -> None:
        """Добавление информации о разработке"""
        text = "Настоящий стандарт разработан и внесен (наименование разработчика)"
        self.doc.add_paragraph(text)

    def _add_approval_info(self) -> None:
        """Добавление информации об утверждении"""
        text = "Утвержден и введен в действие (дата введения в действие)"
        self.doc.add_paragraph(text)

    def _add_replacement_info(self) -> None:
        """Добавление информации о замене"""
        replaces = self.config['document']['compliance'].get('replaces', '')
        text = f"Настоящий стандарт заменяет {replaces}" if replaces else ""
        if text:
            self.doc.add_paragraph(text)

    def _add_patent_notice(self) -> None:
        """Добавление патентного уведомления"""
        text = "Сведения о действии стандарта и патентные ограничения приведены в приложении"
        self.doc.add_paragraph(text)