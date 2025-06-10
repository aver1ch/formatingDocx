from docx import Document
from typing import Dict, Any, Optional
import os
import logging
from datetime import datetime

class TemplateManager:
    """Менеджер для работы с шаблонами документов"""
    
    def __init__(self, templates_dir: str = "doc_editor/templates"):
        self.templates_dir = templates_dir
        self.logger = logging.getLogger(f"{__name__}.{self.__class__.__name__}")
    
    def load_template(self, template_name: str) -> Optional[Document]:
        """
        Загружает шаблон документа
        
        :param template_name: Имя шаблона без расширения
        :return: Document объект или None при ошибке
        """
        template_path = os.path.join(self.templates_dir, f"{template_name}.docx")
        
        if not os.path.exists(template_path):
            self.logger.error(f"Шаблон не найден: {template_path}")
            return None
            
        try:
            return Document(template_path)
        except Exception as e:
            self.logger.error(f"Ошибка загрузки шаблона {template_path}: {e}")
            return None
    
    def replace_placeholders(self, template_doc: Document, data: Dict[str, str]) -> Document:
        """
        Заменяет плейсхолдеры в шаблоне на реальные данные
        
        :param template_doc: Документ шаблона
        :param data: Словарь замен {placeholder: value}
        :return: Обновленный документ
        """
        try:
            # Замена в параграфах
            for paragraph in template_doc.paragraphs:
                self._replace_in_paragraph(paragraph, data)
            
            # Замена в таблицах
            for table in template_doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_in_paragraph(paragraph, data)
            
            self.logger.info(f"Заменено {len(data)} плейсхолдеров в шаблоне")
            return template_doc
            
        except Exception as e:
            self.logger.error(f"Ошибка замены плейсхолдеров: {e}")
            return template_doc
    
    def _replace_in_paragraph(self, paragraph, data: Dict[str, str]) -> None:
        """Замена плейсхолдеров в параграфе"""
        for placeholder, value in data.items():
            placeholder_text = f"{{{{{placeholder}}}}}"
            if placeholder_text in paragraph.text:
                # Проходим по всем runs и заменяем
                for run in paragraph.runs:
                    if placeholder_text in run.text:
                        run.text = run.text.replace(placeholder_text, str(value))
    
    def insert_template_content(self, target_doc: Document, template_doc: Document) -> None:
        """
        Вставляет содержимое шаблона в начало целевого документа
        
        :param target_doc: Целевой документ
        :param template_doc: Документ шаблона
        """
        try:
            # Собираем существующие параграфы
            existing_paragraphs = []
            for paragraph in target_doc.paragraphs:
                existing_paragraphs.append(paragraph.text)
            
            # Очищаем документ
            for paragraph in target_doc.paragraphs:
                p_element = paragraph._element
                p_element.getparent().remove(p_element)
            
            # Добавляем содержимое шаблона
            for paragraph in template_doc.paragraphs:
                new_paragraph = target_doc.add_paragraph()
                new_paragraph.text = paragraph.text
                # Копируем форматирование
                if paragraph.style:
                    try:
                        new_paragraph.style = paragraph.style
                    except:
                        pass  # Если стиль не найден, используем базовый
            
            # Добавляем разрыв страницы
            target_doc.add_page_break()
            
            # Восстанавливаем существующее содержимое
            for text in existing_paragraphs:
                if text.strip():  # Добавляем только непустые параграфы
                    target_doc.add_paragraph(text)
            
            self.logger.info("Содержимое шаблона успешно вставлено в начало документа")
            
        except Exception as e:
            self.logger.error(f"Ошибка вставки содержимого шаблона: {e}")
            raise

