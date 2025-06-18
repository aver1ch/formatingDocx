# title_page.py
from docxtpl import DocxTemplate, InlineImage
from docxcompose.composer import Composer
from docx.shared import Mm
from docx import Document
import logging
from typing import Dict, Any
import os

logger = logging.getLogger(__name__)

class DocumentFormattingError(Exception):
    """Базовое исключение для ошибок форматирования документа."""
    pass

class TitlePageManager:
    def __init__(self, config: Dict[str, Any]):
        self.config = config
        self.logger = logger

    def add_title_page(self, source_doc_path: str, output_path: str) -> None:
        """
        Добавляет титульный лист к документу.
        
        Args:
            source_doc_path: Путь к исходному документу
            output_path: Путь для сохранения результата
            
        Raises:
            DocumentFormattingError: Если не удалось добавить титульный лист
        """
        try:
            title_config = self.config['document']['structure'].get('title_page', {})
            if not title_config:
                self.logger.warning("Конфиг титульного листа не найден, пропускаем добавление")
                return
            print(title_config)
            elements = {}
            for item in title_config.get('elements', []):
                if item:  # Проверяем, что словарь не пустой
                    key, value = next(iter(item.items()))  # Берем первый (и единственный) элемент
                    elements[key] = value

            # Рендерим титульный лист
            title_doc = DocxTemplate(title_config['template_path'])
            context = {
                'agency_name': elements.get('agency_name', ''),
                'standart_type': elements.get('standart_type', ''),
                'image': InlineImage(title_doc, title_config['image_path'], 
                                   width=Mm(42)),
                'designation': elements.get('designation', ''),
                'title': elements.get('title', ''),
                'status': elements.get('status', ''),
                'city': elements.get('city', ''),
                'publisher_info': elements.get('publisher_info', ''),
                'current_year': elements.get('current_year', ''),
            }
            title_doc.render(context)
            title_doc.save("temp_title.docx")

            # Объединяем документы
            composer = Composer(Document())
            composer.append(Document("temp_title.docx"))
            composer.append(Document(source_doc_path))
            composer.save(output_path)
            os.remove("temp_title.docx")
            self.logger.info("Титульный лист успешно добавлен")
        except Exception as e:
            self.logger.error(f"Ошибка добавления титульного листа: {e}")
            raise DocumentFormattingError(f"Ошибка добавления титульного листа: {e}")