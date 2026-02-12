import logging
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.shared import Pt
from typing import Optional, List, Dict, Any

from ..models import DocumentConfig, ProcessorError, HeaderTextPart

logger = logging.getLogger(__name__)


class HeaderFooterProcessor:
    """Обработчик колонтитулов документа."""

    def __init__(self, doc: Document, config: DocumentConfig):
        """
        Инициализация процессора колонтитулов.

        Args:
            doc: Объект документа.
            config: Конфигурация (типизированный объект).
        """
        self.doc = doc
        self.config = config
        self.logger = logger

    def apply(self) -> None:
        """Применяет настройки колонтитулов из конфигурации."""
        try:
            headers_config = self.config.structure.numbering.headers
            if not headers_config.enabled:
                self.logger.info("Колонтитулы отключены в конфигурации")
                return

            self.logger.info("Начало применения колонтитулов")
            self._apply_headers_footers()
            self.logger.info("Колонтитулы успешно применены")
        except Exception as e:
            self.logger.error(f"Ошибка применения колонтитулов: {e}")
            raise ProcessorError(f"Ошибка применения колонтитулов: {e}")

    def _apply_headers_footers(self) -> None:
        """Применяет колонтитулы к документу."""
        headers_config = self.config.structure.numbering.headers
        logger.debug(f"Headers config: right_parts={len(headers_config.right_parts)}, left_parts={len(headers_config.left_parts)}")
        
        self.doc.settings.odd_and_even_pages_header_footer = True

        for i, section in enumerate(self.doc.sections):
            section.different_first_page_header_footer = True

            # Очищаем колонтитул первой страницы (титульной)
            self._clear_element(section.first_page_header)
            self._clear_element(section.first_page_footer)

            # Нечетные страницы: справа
            # Используем right_parts если они есть, иначе fallback на left (строка)
            if headers_config.right_parts:
                logger.debug(f"Section {i}: adding right_parts to header")
                self._add_text_parts_to_element(
                    section.header,
                    headers_config.right_parts,
                    'right'
                )
            else:
                logger.debug(f"Section {i}: adding left string to header")
                self._add_text_to_element(
                    section.header,
                    headers_config.left,
                    'right'
                )

            # Четные страницы: слева
            # Используем left_parts если они есть, иначе fallback на right (строка)
            if headers_config.left_parts:
                logger.debug(f"Section {i}: adding left_parts to even header")
                self._add_text_parts_to_element(
                    section.even_page_header,
                    headers_config.left_parts,
                    'left'
                )
            else:
                logger.debug(f"Section {i}: adding right string to even header")
                self._add_text_to_element(
                    section.even_page_header,
                    headers_config.right,
                    'left'
                )

            # Нумерация страниц
            if headers_config.page_numbers:
                self._add_page_number(section.footer, 'right')
                self._add_page_number(section.even_page_footer, 'left')

    def _clear_element(self, element) -> None:
        """Очищает содержимое колонтитула."""
        for paragraph in element.paragraphs:
            paragraph.clear()
        if not element.paragraphs:
            element.add_paragraph()

    def _add_text_parts_to_element(self, element, text_parts: List[HeaderTextPart], align: Optional[str] = None) -> None:
        """Добавляет текст с поддержкой форматирования (жирный, курсив и т.д.)."""
        logger.debug(f"Adding {len(text_parts)} text parts to element with align={align}")
        
        if element.paragraphs:
            paragraph = element.paragraphs[0]
            paragraph.clear()
        else:
            paragraph = element.add_paragraph()

        # Устанавливаем выравнивание
        if align == 'left':
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif align == 'right':
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        elif align == 'center':
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Добавляем каждую часть с её форматированием
        main_family = self.config.general.fonts['main'].get('family', 'Arial')
        
        for part in text_parts:
            logger.debug(f"Adding part: {part.text!r}, bold={part.bold}")
            run = paragraph.add_run(part.text)
            
            # Применяем форматирование
            run.bold = part.bold
            run.italic = part.italic
            
            # Применяем шрифт
            font_family = part.font_family or main_family
            run.font.name = font_family
            
            # Также установим рFonts на уровне XML для надежности
            try:
                rPr = run._element.find(qn('w:rPr'))
                if rPr is None:
                    rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
                    run._element.insert(0, rPr)
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}></w:rFonts>')
                    rPr.append(rFonts)
                rFonts.set(qn('w:ascii'), font_family)
                rFonts.set(qn('w:hAnsi'), font_family)
                rFonts.set(qn('w:cs'), font_family)
            except Exception as e:
                logger.warning(f"Failed to set XML formatting: {e}")

    def _add_text_to_element(self, element, text: str, align: Optional[str] = None) -> None:
        """Добавляет текст в элемент с указанным выравниванием."""
        if element.paragraphs:
            paragraph = element.paragraphs[0]
            paragraph.clear()
        else:
            paragraph = element.add_paragraph()

        # Устанавливаем выравнивание
        if align == 'left':
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif align == 'right':
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        elif align == 'center':
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Добавляем текст и применяем семейство шрифта из конфигурации
        run = paragraph.add_run(text)
        try:
            main_family = self.config.general.fonts['main'].get('family', None)
            if main_family:
                run.font.name = main_family
                # also set run-level rFonts to ensure Word uses the family (override theme)
                rPr = run._element.find(qn('w:rPr'))
                if rPr is None:
                    from docx.oxml import parse_xml
                    rPr = parse_xml(f'<w:rPr {"" if False else nsdecls("w")}></w:rPr>')
                    run._element.insert(0, rPr)
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    from docx.oxml import parse_xml
                    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}></w:rFonts>')
                    rPr.append(rFonts)
                rFonts.set(qn('w:ascii'), main_family)
                rFonts.set(qn('w:hAnsi'), main_family)
                rFonts.set(qn('w:cs'), main_family)
        except Exception:
            # not critical — continue on error
            pass

    def _add_page_number(self, footer, align: str = 'center') -> None:
        """Добавляет номер страницы в футер."""
        if footer.paragraphs:
            paragraph = footer.paragraphs[0]
            paragraph.clear()
        else:
            paragraph = footer.add_paragraph()

        # Устанавливаем выравнивание
        if align == 'left':
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif align == 'right':
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        else:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Добавляем поле номера страницы
        run = paragraph.add_run()
        fld_char_begin = OxmlElement('w:fldChar')
        fld_char_begin.set(qn('w:fldCharType'), 'begin')

        instr_text = OxmlElement('w:instrText')
        instr_text.text = 'PAGE'

        fld_char_end = OxmlElement('w:fldChar')
        fld_char_end.set(qn('w:fldCharType'), 'end')

        run._element.append(fld_char_begin)
        run._element.append(instr_text)
        run._element.append(fld_char_end)
        # Применим основной шрифт к полю номера страницы, если задан
        try:
            main_family = self.config.general.fonts['main'].get('family', None)
            if main_family:
                run.font.name = main_family
                rPr = run._element.find(qn('w:rPr'))
                if rPr is None:
                    from docx.oxml import parse_xml
                    rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
                    run._element.insert(0, rPr)
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    from docx.oxml import parse_xml
                    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}></w:rFonts>')
                    rPr.append(rFonts)
                rFonts.set(qn('w:ascii'), main_family)
                rFonts.set(qn('w:hAnsi'), main_family)
                rFonts.set(qn('w:cs'), main_family)
        except Exception:
            pass
