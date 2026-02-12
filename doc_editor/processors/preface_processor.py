"""
PrefaceProcessor - Компонент для добавления предисловия/преамбулы в документ.

Функционал:
- Добавление содержания предисловия после титула и оглавления
- Правильное позиционирование в документе
- Форматирование и управление через конфигурацию
- Поддержка многострочного содержимого
"""

import logging
from typing import Optional
from docx.document import Document

logger = logging.getLogger(__name__)


class PrefaceProcessor:
    """
    Обработчик предисловия документа - Этап 3 Фазы 2.
    
    Добавляет содержание предисловия в документ после титула и оглавления,
    перед основным содержимым.
    """
    
    def __init__(self, config):
        """
        Инициализация процессора предисловия.
        
        Args:
            config: DocumentConfig с настройками предисловия
        """
        self.config = config
        self.logger = logger
    
    def add_preface(self, document: Document) -> None:
        """
        Добавить предисловие в документ.
        
        Предисловие добавляется:
        1. После титула (если есть)
        2. После оглавления (если есть)
        3. Перед основным содержимым
        
        Args:
            document: Документ Word для обработки
            
        Returns:
            None
        """
        if not self.config.structure.document_structure.preface.enabled:
            self.logger.debug("Предисловие отключено в конфигурации")
            return
        
        self.logger.info("Начинаю добавление предисловия в документ")
        
        try:
            preface_content = self.config.structure.document_structure.preface.content
            
            if not preface_content or not preface_content.strip():
                self.logger.warning("Содержание предисловия пусто")
                return
            
            # Вставить предисловие в документ
            self._insert_preface_to_document(document, preface_content)
            
            self.logger.info("Предисловие успешно добавлено в документ")
            
        except Exception as e:
            self.logger.error(f"Ошибка при добавлении предисловия: {str(e)}", exc_info=True)
            raise
    
    def _insert_preface_to_document(self, document: Document, content: str) -> None:
        """
        Вставить содержание предисловия в документ.
        
        Вставляет предисловие в начало документа (или после существующего содержимого),
        разбивая содержание на параграфы по переносам строк.
        
        Args:
            document: Документ для вставки
            content: Содержание предисловия
            
        Returns:
            None
        """
        if not document.paragraphs:
            self.logger.debug("Документ пуст, добавляю предисловие в начало")
            insert_index = 0
        else:
            # Вставить перед первым параграфом
            insert_index = 0
        
        # Разбить содержание на строки
        preface_lines = content.split('\n')
        
        # Вставить содержание предисловия
        for i, line in enumerate(preface_lines):
            line = line.strip()
            
            if not line:
                # Пустая строка - пропустить
                continue
            
            # Вставить параграф
            if document.paragraphs and insert_index < len(document.paragraphs):
                preface_para = document.paragraphs[insert_index].insert_paragraph_before(line)
            else:
                preface_para = document.add_paragraph(line)
            
            # Установить стиль (Normal)
            preface_para.style = 'Normal'
        
        self.logger.debug(f"Вставлено {len(preface_lines)} строк предисловия")
    
    def _create_preface_paragraph(self, content: str, style: str = 'Normal'):
        """
        Создать параграф для предисловия.
        
        Args:
            content: Содержание параграфа
            style: Стиль параграфа (по умолчанию Normal)
            
        Returns:
            Объект параграфа (для встраивания в документ)
        """
        self.logger.debug(f"Создание параграфа предисловия со стилем '{style}'")
        
        # Это вспомогательный метод, детали реализации зависят от python-docx
        return {
            'text': content,
            'style': style
        }

