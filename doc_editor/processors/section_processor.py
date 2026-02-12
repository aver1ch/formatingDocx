"""
SectionProcessor - Компонент для многоуровневой нумерации разделов документа.
Реализует автоматическую нумерацию заголовков согласно ГОСТ Р 1.5-2004.
"""

import logging
from typing import List, Optional
from docx.document import Document
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


class SectionProcessor:
    """Обработчик многоуровневой нумерации разделов."""
    
    # Соответствие стилей заголовков уровням
    HEADING_LEVELS = {
        'Heading 1': 0,
        'Heading 2': 1,
        'Heading 3': 2,
    }
    
    def __init__(self, config):
        """
        Инициализация процессора.
        
        Args:
            config: DocumentConfig с настройками нумерации
        """
        self.config = config
        self.section_numbers: List[int] = [0, 0, 0]  # Счетчики для каждого уровня
        self.logger = logger
    
    def apply_section_numbering(self, document: Document) -> None:
        """
        Применить многоуровневую нумерацию разделов к документу.
        
        Args:
            document: Документ Word для обработки
        """
        # Проверяем, включена ли нумерация разделов
        if not self.config.structure.document_structure.sections.enabled:
            self.logger.debug("Нумерация разделов отключена в конфигурации")
            return
        
        self.logger.info("Начинаю применение многоуровневой нумерации разделов...")
        self.reset_numbering()
        
        processed_count = 0
        for paragraph in document.paragraphs:
            style_name = paragraph.style.name
            
            # Проверяем, является ли это заголовком
            if style_name not in self.HEADING_LEVELS:
                continue
            
            level = self.HEADING_LEVELS[style_name]
            self._process_heading(paragraph, level)
            processed_count += 1
        
        self.logger.info(f"Обработано {processed_count} заголовков с нумерацией")
    
    def _process_heading(self, paragraph: Paragraph, level: int) -> None:
        """
        Обработать один заголовок: обновить номер раздела.
        
        Args:
            paragraph: Абзац с заголовком
            level: Уровень заголовка (0, 1, 2)
        """
        # Сохраняем текущий текст
        current_text = paragraph.text.strip()
        
        # Проверяем, не начинается ли уже с цифры (уже нумерован)
        if current_text and current_text[0].isdigit():
            # Парсим существующий номер и обновляем счетчики
            self._parse_and_update_from_existing(current_text, level)
            return
        
        # Обновляем счетчик для текущего уровня
        self._update_section_number(level)
        
        # Получаем отформатированный номер
        section_num = self._get_section_number(level)
        
        # Очищаем параграф
        for run in paragraph.runs:
            run.clear()
        
        # Добавляем номер раздела
        num_run = paragraph.add_run(f"{section_num} ")
        num_run.bold = True
        
        # Добавляем оригинальный текст
        text_run = paragraph.add_run(current_text)
        text_run.bold = True
        
        self.logger.debug(f"Добавлена нумерация '{section_num}' к заголовку")
    
    def _update_section_number(self, level: int) -> None:
        """
        Обновить счетчик для указанного уровня.
        Увеличивает счетчик текущего уровня и сбрасывает все подуровни.
        
        Args:
            level: Уровень (0, 1, 2)
        """
        # Увеличиваем счетчик текущего уровня
        self.section_numbers[level] += 1
        
        # Сбрасываем все подуровни
        for i in range(level + 1, len(self.section_numbers)):
            self.section_numbers[i] = 0
    
    def _get_section_number(self, level: int) -> str:
        """
        Получить отформатированный номер раздела.
        Например: "1", "1.1", "1.1.1"
        
        Args:
            level: Уровень (0, 1, 2)
        
        Returns:
            Строка с номером раздела
        """
        numbers = self.section_numbers[:level + 1]
        return ".".join(str(n) for n in numbers)
    
    def _parse_and_update_from_existing(self, text: str, level: int) -> None:
        """
        Парсить существующий номер раздела и обновить счетчики.
        Используется для обработки уже нумерованных заголовков.
        
        Args:
            text: Текст абзаца, начинающийся с номера
            level: Уровень (0, 1, 2)
        """
        # Найти конец номера (до пробела)
        space_idx = text.find(" ")
        if space_idx == -1:
            return
        
        number_str = text[:space_idx]
        
        try:
            # Парсим номер (1, 1.1, 1.1.1)
            parts = number_str.split(".")
            for i, part in enumerate(parts[:level + 1]):
                self.section_numbers[i] = int(part)
            
            # Сбрасываем подуровни
            for i in range(level + 1, len(self.section_numbers)):
                self.section_numbers[i] = 0
                
        except (ValueError, IndexError):
            # Если парсинг не удался, просто логируем
            self.logger.warning(f"Не удалось спарсить номер: {number_str}")
    
    def reset_numbering(self) -> None:
        """Сбросить счетчики нумерации на начальные значения."""
        self.section_numbers = [0, 0, 0]
        self.logger.debug("Счетчики нумерации сброшены")
    
    def get_current_section_number(self) -> str:
        """
        Получить текущий номер раздела.
        Полезно для встраивания в другие компоненты.
        
        Returns:
            Строка с текущим номером раздела
        """
        return ".".join(str(n) for n in self.section_numbers if n > 0 or self.section_numbers.index(n) == 0)
