import logging
import os
from docxtpl import DocxTemplate, InlineImage
from docxcompose.composer import Composer
from docx import Document
from docx.shared import Mm, Pt
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from typing import Dict, Any

from ..models import DocumentConfig, ProcessorError

logger = logging.getLogger(__name__)


class TitleProcessor:
    """Обработчик титульного листа документа."""

    def __init__(self, config: DocumentConfig):
        """
        Инициализация процессора титула.

        Args:
            config: Конфигурация (типизированный объект).
        """
        self.config = config
        self.logger = logger

    def apply(self, source_doc_path: str, output_path: str) -> None:
        """
        Добавляет титульный лист к документу.

        Args:
            source_doc_path: Путь к исходному документу.
            output_path: Путь для сохранения результата.

        Raises:
            ProcessorError: Если не удалось добавить титульный лист.
        """
        try:
            title_config = self.config.structure.title_page
            if not title_config.enabled:
                self.logger.info("Титульный лист отключен в конфигурации")
                return

            self.logger.info("Начало добавления титульного листа")
            self._add_title_page(source_doc_path, output_path, title_config)
            self.logger.info("Титульный лист успешно добавлен")
        except Exception as e:
            self.logger.error(f"Ошибка добавления титульного листа: {e}")
            raise ProcessorError(f"Ошибка добавления титульного листа: {e}")

    def _add_title_page(self, source_doc_path: str, output_path: str, title_config: Any) -> None:
        """Добавляет титульный лист к документу."""
        # Парсим элементы конфигурации в словарь
        elements = self._parse_elements(title_config.elements)

        # Рендерим титульный лист
        title_doc = DocxTemplate(title_config.template_path)
        context = {
            'agency_name': elements.get('agency_name', ''),
            'st_type': elements.get('standart_type', ''),
            'image': InlineImage(title_doc, title_config.image_path, width=Mm(42)),
            'designation': elements.get('designation', ''),
            'title': elements.get('title', ''),
            'status': elements.get('status', ''),
            'city': elements.get('city', ''),
            'publisher_info': elements.get('publisher_info', ''),
            'current_year': elements.get('current_year', ''),
        }
        title_doc.render(context)
        # После рендера принудительно установим семейство шрифта из конфигурации
        try:
            main_family = self.config.general.fonts['main'].get('family', None)
            if main_family:
                # применяем к параграфам и таблицам рендеренного титула
                self._apply_font_to_doc(title_doc, main_family)
                # если в шаблоне есть стиль Custom_Title — установим его font.name тоже
                try:
                    st = title_doc.styles['Custom_Title']
                    st.font.name = main_family
                except Exception:
                    pass
        except Exception:
            pass

        title_doc.save("temp_title.docx")

        # Применяем дополнительное форматирование (spacing, table formatting)
        self._apply_formatting_to_doc(title_doc, title_config)

        title_doc.save("temp_title.docx")

        # Объединяем документы
        composer = Composer(Document())
        composer.append(Document("temp_title.docx"))
        composer.append(Document(source_doc_path))
        composer.save(output_path)
        
        # Очищаем временный файл
        os.remove("temp_title.docx")

    @staticmethod
    def _parse_elements(elements_list: list) -> Dict[str, str]:
        """Парсит список элементов конфигурации в словарь."""
        result = {}
        for item in elements_list:
            if item and isinstance(item, dict):
                key, value = next(iter(item.items()))
                result[key] = value
        return result

    def _apply_font_to_doc(self, doc: Document, family: str) -> None:
        """Apply run-level font family to all runs in a Document (paragraphs and table cells)."""
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                try:
                    run.font.name = family
                    # ensure run rFonts are set to override theme
                    rPr = run._element.find(qn('w:rPr'))
                    if rPr is None:
                        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
                        run._element.insert(0, rPr)
                    rFonts = rPr.find(qn('w:rFonts'))
                    if rFonts is None:
                        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}></w:rFonts>')
                        rPr.append(rFonts)
                    rFonts.set(qn('w:ascii'), family)
                    rFonts.set(qn('w:hAnsi'), family)
                    rFonts.set(qn('w:cs'), family)
                except Exception:
                    pass

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            try:
                                run.font.name = family
                                rPr = run._element.find(qn('w:rPr'))
                                if rPr is None:
                                    rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
                                    run._element.insert(0, rPr)
                                rFonts = rPr.find(qn('w:rFonts'))
                                if rFonts is None:
                                    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}></w:rFonts>')
                                    rPr.append(rFonts)
                                rFonts.set(qn('w:ascii'), family)
                                rFonts.set(qn('w:hAnsi'), family)
                                rFonts.set(qn('w:cs'), family)
                            except Exception:
                                pass
    def _apply_formatting_to_doc(self, doc: Document, title_config: Any) -> None:
        """
        Применяет форматирование (spacing, table formatting) к документу.
        
        Args:
            doc: Документ для форматирования.
            title_config: Конфигурация титульной страницы.
        """
        # Применяем line spacing к параграфам
        if title_config.line_spacing:
            self._apply_line_spacing_to_doc(doc, title_config.line_spacing)
        
        # Применяем spacing_before и spacing_after к параграфам
        if title_config.spacing_before or title_config.spacing_after:
            self._apply_paragraph_spacing(doc, title_config.spacing_before, title_config.spacing_after)
        
        # Применяем форматирование таблиц если необходимо
        if title_config.table_format and title_config.table_format.apply_font:
            main_family = self.config.general.fonts['main'].get('family', None)
            if main_family:
                self._format_tables(doc, main_family, title_config.table_format)
        
        self.logger.debug("Форматирование документа успешно применено")

    def _apply_line_spacing_to_doc(self, doc: Document, line_spacing: float) -> None:
        """
        Применяет межстрочный интервал ко всем параграфам в документе.
        
        Args:
            doc: Документ.
            line_spacing: Коэффициент межстрочного интервала (e.g., 1.5 для 1.5 интервала).
        """
        for paragraph in doc.paragraphs:
            paragraph.paragraph_format.line_spacing = line_spacing
        
        # Также применяем к параграфам в таблицах
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.line_spacing = line_spacing
        
        self.logger.debug(f"Межстрочный интервал {line_spacing} применен ко всему документу")

    def _apply_paragraph_spacing(self, doc: Document, spacing_before: float, spacing_after: float) -> None:
        """
        Применяет spacing before/after ко всем параграфам.
        
        Args:
            doc: Документ.
            spacing_before: Промежуток перед параграфом в пунктах (pt).
            spacing_after: Промежуток после параграфа в пунктах (pt).
        """
        for paragraph in doc.paragraphs:
            if spacing_before > 0:
                paragraph.paragraph_format.space_before = Pt(spacing_before)
            if spacing_after > 0:
                paragraph.paragraph_format.space_after = Pt(spacing_after)
        
        # Также применяем к параграфам в таблицах
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if spacing_before > 0:
                            paragraph.paragraph_format.space_before = Pt(spacing_before)
                        if spacing_after > 0:
                            paragraph.paragraph_format.space_after = Pt(spacing_after)
        
        self.logger.debug(f"Spacing before={spacing_before}, after={spacing_after} применены")

    def _format_tables(self, doc: Document, font_family: str, table_format: Any) -> None:
        """
        Применяет форматирование к таблицам.
        
        Args:
            doc: Документ.
            font_family: Семейство шрифтов.
            table_format: Конфигурация форматирования таблиц.
        """
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Применяем форматирование шрифта к ячейкам
                    if table_format.apply_font:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                try:
                                    run.font.name = font_family
                                    # Установим XML-уровневые атрибуты
                                    rPr = run._element.find(qn('w:rPr'))
                                    if rPr is None:
                                        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
                                        run._element.insert(0, rPr)
                                    rFonts = rPr.find(qn('w:rFonts'))
                                    if rFonts is None:
                                        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}></w:rFonts>')
                                        rPr.append(rFonts)
                                    rFonts.set(qn('w:ascii'), font_family)
                                    rFonts.set(qn('w:hAnsi'), font_family)
                                    rFonts.set(qn('w:cs'), font_family)
                                except Exception:
                                    pass
                    
                    # Применяем spacing если задано
                    if table_format.apply_spacing:
                        for paragraph in cell.paragraphs:
                            if hasattr(table_format, 'line_spacing') and table_format.line_spacing:
                                paragraph.paragraph_format.line_spacing = table_format.line_spacing
        
        self.logger.debug("Форматирование таблиц успешно применено")