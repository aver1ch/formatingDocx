"""
TOCProcessor - Компонент для автоматического построения оглавления.

Функционал:
- Извлечение всех заголовков (Heading 1, 2, 3) из документа
- Определение уровня заголовка и расчет номера страницы
- Форматирование оглавления с правильными отступами
- Вставка оглавления в начало документа (после титула, перед основным текстом)
- Поддержка конфигурации: включение/отключение, кастомный заголовок, номера страниц, уровни
"""

import logging
from typing import List, Dict, Optional
from docx.document import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)


class TOCProcessor:
    """
    Обработчик оглавления (Table of Contents) - Этап 2 Фазы 2.
    
    Автоматически создает и вставляет оглавление документа на основе заголовков.
    """
    
    # Маппинг стилей заголовков на уровни (0 = Level 1, 1 = Level 2, 2 = Level 3)
    HEADING_STYLES = {
        'Heading 1': 0,
        'Heading 2': 1,
        'Heading 3': 2,
    }
    
    def __init__(self, config):
        """
        Инициализация процессора оглавления.
        
        Args:
            config: DocumentConfig с настройками оглавления
        """
        self.config = config
        self.logger = logger
    
    def create_toc(self, document: Document) -> None:
        """
        Создать и вставить оглавление в начало документа.
        
        Процесс:
        1. Проверить, включено ли оглавление в конфигурации
        2. Извлечь все заголовки из документа
        3. Построить записи оглавления с номерами страниц
        4. Отформатировать записи с правильными отступами
        5. Вставить оглавление в начало документа
        
        Args:
            document: Документ Word для обработки
            
        Returns:
            None
        """
        if not self.config.structure.document_structure.toc.enabled:
            self.logger.debug("Оглавление отключено в конфигурации")
            return
        
        self.logger.info("Начинаю создание оглавления")
        
        try:
            # Шаг 1: Извлечь все заголовки
            headings = self._extract_headings(document)
            
            if not headings:
                self.logger.warning("В документе не найдены заголовки для оглавления")
                return
            
            # Шаг 2: Построить записи оглавления
            toc_entries = self._build_toc_entries(document, headings)
            
            # Шаг 3: Построить строки оглавления
            toc_lines = self._build_toc_lines(toc_entries)
            
            # Шаг 4: Вставить в документ
            self._insert_toc_to_document(document, toc_lines)
            
            self.logger.info(f"Оглавление успешно создано ({len(toc_entries)} записей)")
            
        except Exception as e:
            self.logger.error(f"Ошибка при создании оглавления: {str(e)}", exc_info=True)
            raise
    
    def _extract_headings(self, document: Document) -> List:
        """
        Извлечь все заголовки из документа.
        
        Извлекает параграфы со стилями Heading 1, 2, 3 в порядке их появления.
        
        Args:
            document: Исходный документ
            
        Returns:
            Список объектов Paragraph, содержащих заголовки
        """
        headings = []
        
        for paragraph in document.paragraphs:
            if paragraph.style.name in self.HEADING_STYLES:
                headings.append(paragraph)
        
        self.logger.debug(f"Найдено заголовков: {len(headings)}")
        return headings
    
    def _get_heading_level(self, style_name: str) -> int:
        """
        Определить уровень заголовка по названию стиля.
        
        Args:
            style_name: Название стиля (например, "Heading 1")
            
        Returns:
            Уровень заголовка: 0 для Heading 1, 1 для Heading 2, 2 для Heading 3
        """
        return self.HEADING_STYLES.get(style_name, 0)
    
    def _get_paragraph_page_number(self, document: Document, paragraph) -> int:
        """
        Получить номер страницы для параграфа.
        
        Используется метод подсчета на основе среднего количества строк на странице.
        В реальных документах Word номер страницы зависит от множества факторов,
        поэтому используется приблизительный расчет.
        
        Args:
            document: Документ Word
            paragraph: Параграф, для которого нужно найти номер страницы
            
        Returns:
            Номер страницы (1-indexed)
        """
        try:
            # Получить индекс параграфа
            para_index = document.paragraphs.index(paragraph)
            
            # Приблизительно: 55 строк на страницу (зависит от форматирования)
            # Это примерная оценка для стандартного документа А4
            lines_per_page = 55
            page_num = (para_index // lines_per_page) + 1
            
            return max(1, page_num)
            
        except (ValueError, IndexError):
            self.logger.warning(f"Не удалось определить номер страницы для параграфа")
            return 1
    
    def _build_toc_entries(self, document: Document, headings: List) -> List[Dict]:
        """
        Построить записи оглавления с информацией о каждом заголовке.
        
        Для каждого заголовка определяется:
        - Уровень (0, 1 или 2)
        - Текст заголовка
        - Номер страницы
        
        Фильтруются заголовки по максимальному уровню из конфигурации.
        
        Args:
            document: Исходный документ
            headings: Список найденных заголовков
            
        Returns:
            Список словарей с ключами: level, text, page_num
        """
        entries = []
        max_levels = self.config.structure.document_structure.toc.levels
        
        for heading in headings:
            level = self._get_heading_level(heading.style.name)
            
            # Пропустить, если уровень превышает максимальный
            if level >= max_levels:
                continue
            
            page_num = self._get_paragraph_page_number(document, heading)
            
            entries.append({
                'level': level,
                'text': heading.text,
                'page_num': page_num
            })
        
        self.logger.debug(f"Построено записей оглавления: {len(entries)}")
        return entries
    
    def _build_toc_lines(self, entries: List[Dict]) -> List[str]:
        """
        Построить строки оглавления с правильным форматированием.
        
        Каждая строка содержит:
        - Отступ, соответствующий уровню иерархии (2 пробела на уровень)
        - Текст заголовка
        - Многоточие (если включены номера страниц)
        - Номер страницы (если включены в конфигурации)
        
        Пример:
        ```
        ОГЛАВЛЕНИЕ
        1. Introduction...1
          1.1 Background...2
            1.1.1 Details...3
        2. Methods...4
        ```
        
        Args:
            entries: Список записей оглавления
            
        Returns:
            Список строк для вставки в документ
        """
        lines = []
        show_page_numbers = self.config.structure.document_structure.toc.page_numbers
        
        for entry in entries:
            level = entry['level']
            text = entry['text']
            page_num = entry['page_num']
            
            # Создать отступ на основе уровня (2 пробела на уровень)
            indent = "  " * level
            
            # Построить строку
            if show_page_numbers:
                # Формат с номером страницы и многоточием
                line = f"{indent}{text}...{page_num}"
            else:
                # Формат без номера страницы
                line = f"{indent}{text}"
            
            lines.append(line)
        
        self.logger.debug(f"Построено строк оглавления: {len(lines)}")
        return lines
    
    def _insert_toc_to_document(self, document: Document, toc_lines: List[str]) -> None:
        """
        Вставить оглавление в начало документа.
        
        Вставляет:
        1. Заголовок "ОГЛАВЛЕНИЕ"
        2. Пустую строку
        3. Все строки оглавления
        4. Пустую строку (разделитель от основного текста)
        
        Заголовок вставляется со стилем Heading 1.
        Строки оглавления - со стилем Normal.
        
        Args:
            document: Документ для вставки
            toc_lines: Список строк оглавления
            
        Returns:
            None
        """
        toc_title = self.config.structure.document_structure.toc.title
        
        # Вставить в начало документа (перед первым параграфом)
        insert_index = 0
        
        # Если первый параграф пустой, вставить туда
        if document.paragraphs and not document.paragraphs[0].text.strip():
            insert_index = 0
        
        # Вставить заголовок оглавления
        if document.paragraphs and insert_index < len(document.paragraphs):
            title_para = document.paragraphs[insert_index].insert_paragraph_before(toc_title)
        else:
            title_para = document.add_paragraph(toc_title)
        
        # Установить стиль для заголовка
        title_para.style = 'Heading 1'
        self.logger.debug(f"Вставлен заголовок оглавления: '{toc_title}'")
        
        # Вставить пустую строку
        if document.paragraphs and insert_index < len(document.paragraphs):
            spacer1 = document.paragraphs[insert_index].insert_paragraph_before("")
        else:
            spacer1 = document.add_paragraph("")
        
        # Вставить строки оглавления
        for i, line in enumerate(toc_lines):
            if document.paragraphs and insert_index + i + 2 < len(document.paragraphs):
                toc_para = document.paragraphs[insert_index + i + 2].insert_paragraph_before(line)
            else:
                toc_para = document.add_paragraph(line)
            
            # Установить стиль (Normal)
            toc_para.style = 'Normal'
        
        self.logger.debug(f"Вставлено {len(toc_lines)} строк оглавления")
        
        # Вставить разделитель (пустая строка)
        if document.paragraphs and insert_index + len(toc_lines) + 2 < len(document.paragraphs):
            spacer2 = document.paragraphs[insert_index + len(toc_lines) + 2].insert_paragraph_before("")
        else:
            spacer2 = document.add_paragraph("")
        
        self.logger.info(f"Оглавление вставлено в начало документа")
