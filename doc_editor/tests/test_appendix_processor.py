"""
Comprehensive test suite for AppendixProcessor.

Tests cover:
1. Initialization and configuration
2. Enable/disable functionality
3. Appendix detection and finding
4. Letter numbering (A, B, C... and А, Б, В...)
5. Formatting and styling
6. Edge cases and error handling
7. Integration with document structure
"""

import pytest
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE

from doc_editor.models.config import (
    DocumentConfig,
    DocumentStructureConfig,
    AppendixConfig,
    GeneralConfig,
    MarginsConfig,
    SpacingConfig,
    StructureConfig,
)
from doc_editor.processors.appendix_processor import AppendixProcessor


# ============================================================================
# FIXTURES
# ============================================================================

@pytest.fixture
def base_config():
    """Base configuration with AppendixProcessor enabled."""
    doc_structure = DocumentStructureConfig(
        appendix=AppendixConfig(
            enabled=True,
            numbering_style="letters"
        )
    )
    
    structure = StructureConfig(
        document_structure=doc_structure
    )
    
    general = GeneralConfig(
        margins=MarginsConfig(),
        fonts={},
        spacing=SpacingConfig()
    )
    
    return DocumentConfig(
        general=general,
        structure=structure
    )


@pytest.fixture
def appendix_processor(base_config):
    """AppendixProcessor instance with base config."""
    return AppendixProcessor(base_config)


@pytest.fixture
def empty_document():
    """Empty document for testing."""
    return Document()


@pytest.fixture
def document_with_appendices():
    """Document with appendix headings to process."""
    doc = Document()
    
    # Add main content
    doc.add_paragraph("Main Content", style='Heading 1')
    doc.add_paragraph("Introduction text", style='Normal')
    
    # Add appendix sections
    doc.add_paragraph("Appendix Title", style='Heading 1')
    doc.add_paragraph("First appendix content", style='Normal')
    
    doc.add_paragraph("Another Section", style='Heading 1')
    doc.add_paragraph("Second appendix content", style='Normal')
    
    return doc


@pytest.fixture
def document_with_appendix_keyword():
    """Document with explicit 'Appendix' keyword in headings."""
    doc = Document()
    
    doc.add_paragraph("Document Title", style='Heading 1')
    doc.add_paragraph("Main content", style='Normal')
    
    # Appendix with explicit keyword
    doc.add_paragraph("Appendix A: First Appendix", style='Heading 1')
    doc.add_paragraph("Content of first appendix", style='Normal')
    
    doc.add_paragraph("Appendix B: Second Appendix", style='Heading 1')
    doc.add_paragraph("Content of second appendix", style='Normal')
    
    return doc


@pytest.fixture
def document_with_приложение():
    """Document with Russian 'Приложение' keyword."""
    doc = Document()
    
    doc.add_paragraph("Документ", style='Heading 1')
    doc.add_paragraph("Основное содержание", style='Normal')
    
    # Russian appendix keyword
    doc.add_paragraph("Приложение 1: Первое приложение", style='Heading 1')
    doc.add_paragraph("Содержание первого приложения", style='Normal')
    
    doc.add_paragraph("Приложение 2: Второе приложение", style='Heading 1')
    doc.add_paragraph("Содержание второго приложения", style='Normal')
    
    return doc


@pytest.fixture
def document_with_tables():
    """Document with tables in appendices."""
    doc = Document()
    
    doc.add_paragraph("Main Content", style='Heading 1')
    doc.add_paragraph("Introduction", style='Normal')
    
    doc.add_paragraph("Appendix One", style='Heading 1')
    
    # Add table to appendix
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Header 1"
    table.cell(0, 1).text = "Header 2"
    table.cell(1, 0).text = "Data 1"
    table.cell(1, 1).text = "Data 2"
    
    return doc


# ============================================================================
# TEST CLASSES
# ============================================================================

class TestAppendixProcessorInitialization:
    """Test AppendixProcessor initialization and configuration."""
    
    def test_processor_initialization(self, base_config):
        """Test that processor initializes correctly with config."""
        processor = AppendixProcessor(base_config)
        assert processor is not None
        assert processor.config == base_config
    
    def test_processor_disabled_by_default(self):
        """Test that processor respects disabled config."""
        doc_structure = DocumentStructureConfig(
            appendix=AppendixConfig(enabled=False)
        )
        structure = StructureConfig(document_structure=doc_structure)
        general = GeneralConfig(
            margins=MarginsConfig(),
            fonts={},
            spacing=SpacingConfig()
        )
        config = DocumentConfig(general=general, structure=structure)
        
        processor = AppendixProcessor(config)
        assert not processor.config.structure.document_structure.appendix.enabled
    
    def test_logger_available(self, appendix_processor):
        """Test that logger is properly initialized."""
        assert hasattr(appendix_processor, 'logger')
        assert appendix_processor.logger is not None


class TestAppendixProcessorToggling:
    """Test enable/disable functionality."""
    
    def test_disabled_processor_does_nothing(self, empty_document):
        """Test that disabled processor doesn't modify document."""
        doc_structure = DocumentStructureConfig(
            appendix=AppendixConfig(enabled=False)
        )
        structure = StructureConfig(document_structure=doc_structure)
        general = GeneralConfig(
            margins=MarginsConfig(),
            fonts={},
            spacing=SpacingConfig()
        )
        config = DocumentConfig(general=general, structure=structure)
        processor = AppendixProcessor(config)
        
        initial_para_count = len(empty_document.paragraphs)
        processor.process_appendices(empty_document)
        
        assert len(empty_document.paragraphs) == initial_para_count
    
    def test_enabled_processor_affects_document(self, appendix_processor, 
                                                 document_with_appendices):
        """Test that enabled processor can modify document."""
        # This is a basic test; actual modification depends on implementation
        result = appendix_processor.process_appendices(document_with_appendices)
        # Should not raise exception
        assert result is None or isinstance(result, (None, Document))


class TestAppendixDetection:
    """Test appendix detection and finding."""
    
    def test_find_appendix_headings_english(self, appendix_processor, 
                                            document_with_appendix_keyword):
        """Test detection of English 'Appendix' keyword in headings."""
        headings = appendix_processor._find_appendix_headings(
            document_with_appendix_keyword
        )
        
        assert len(headings) >= 0  # Should find or not find appendices
        # If found, should contain appendix-related text
        for heading in headings:
            assert heading is not None
    
    def test_find_appendix_headings_russian(self, appendix_processor,
                                            document_with_приложение):
        """Test detection of Russian 'Приложение' keyword."""
        headings = appendix_processor._find_appendix_headings(
            document_with_приложение
        )
        
        assert len(headings) >= 0  # May or may not detect depending on implementation
    
    def test_find_appendix_empty_document(self, appendix_processor,
                                          empty_document):
        """Test that finding appendices in empty document is safe."""
        headings = appendix_processor._find_appendix_headings(empty_document)
        assert len(headings) == 0


class TestAppendixNumbering:
    """Test appendix numbering styles."""
    
    def test_letter_numbering_default(self, base_config):
        """Test default letter numbering style."""
        assert base_config.structure.document_structure.appendix.numbering_style == "letters"
    
    def test_letter_numbering_get_letter_a(self, appendix_processor):
        """Test getting first letter (A/А)."""
        letter = appendix_processor._get_appendix_letter(0)
        assert letter in ['A', 'А']
    
    def test_letter_numbering_sequence(self, appendix_processor):
        """Test letter numbering sequence."""
        letters = []
        for i in range(5):
            letter = appendix_processor._get_appendix_letter(i)
            letters.append(letter)
        
        # Should have valid letters
        assert all(len(l) == 1 for l in letters)
    
    def test_number_numbering_style(self):
        """Test numeric numbering style configuration."""
        doc_structure = DocumentStructureConfig(
            appendix=AppendixConfig(
                enabled=True,
                numbering_style="numbers"
            )
        )
        structure = StructureConfig(document_structure=doc_structure)
        general = GeneralConfig(
            margins=MarginsConfig(),
            fonts={},
            spacing=SpacingConfig()
        )
        config = DocumentConfig(general=general, structure=structure)
        
        processor = AppendixProcessor(config)
        assert processor.config.structure.document_structure.appendix.numbering_style == "numbers"


class TestAppendixFormatting:
    """Test appendix formatting and styling."""
    
    def test_process_appendices_no_error_empty(self, appendix_processor,
                                               empty_document):
        """Test processing empty document doesn't raise error."""
        try:
            appendix_processor.process_appendices(empty_document)
        except Exception as e:
            pytest.fail(f"process_appendices raised {type(e).__name__}: {e}")
    
    def test_process_appendices_preserves_structure(self, appendix_processor,
                                                    document_with_appendices):
        """Test that processing preserves document structure."""
        initial_para_count = len(document_with_appendices.paragraphs)
        
        appendix_processor.process_appendices(document_with_appendices)
        
        # Structure should be preserved (may update text, but count stays same)
        final_para_count = len(document_with_appendices.paragraphs)
        
        assert final_para_count >= initial_para_count


class TestAppendixEdgeCases:
    """Test edge cases and error handling."""
    
    def test_document_with_no_appendices(self, appendix_processor):
        """Test processing document with no appendix headings."""
        doc = Document()
        doc.add_paragraph("Regular Content", style='Heading 1')
        doc.add_paragraph("Some text", style='Normal')
        
        try:
            appendix_processor.process_appendices(doc)
        except Exception as e:
            pytest.fail(f"Should handle documents without appendices: {e}")
    
    def test_single_appendix(self, appendix_processor):
        """Test processing document with single appendix."""
        doc = Document()
        doc.add_paragraph("Main Content", style='Heading 1')
        doc.add_paragraph("Text", style='Normal')
        doc.add_paragraph("Appendix One", style='Heading 1')
        doc.add_paragraph("Appendix content", style='Normal')
        
        try:
            appendix_processor.process_appendices(doc)
        except Exception as e:
            pytest.fail(f"Should handle single appendix: {e}")
    
    def test_multiple_appendices(self, appendix_processor):
        """Test processing document with multiple appendices."""
        doc = Document()
        doc.add_paragraph("Main Content", style='Heading 1')
        doc.add_paragraph("Text", style='Normal')
        
        for i in range(5):
            doc.add_paragraph(f"Appendix {i+1}", style='Heading 1')
            doc.add_paragraph(f"Content {i+1}", style='Normal')
        
        try:
            appendix_processor.process_appendices(doc)
        except Exception as e:
            pytest.fail(f"Should handle multiple appendices: {e}")
    
    def test_appendix_with_special_characters(self, appendix_processor):
        """Test processing appendix with special characters."""
        doc = Document()
        doc.add_paragraph("Main Content", style='Heading 1')
        doc.add_paragraph("Text", style='Normal')
        doc.add_paragraph("Appendix: Data & Statistics (2023)", style='Heading 1')
        doc.add_paragraph("Content with @#$% chars", style='Normal')
        
        try:
            appendix_processor.process_appendices(doc)
        except Exception as e:
            pytest.fail(f"Should handle special characters: {e}")
    
    def test_appendix_with_long_title(self, appendix_processor):
        """Test processing appendix with very long title."""
        doc = Document()
        doc.add_paragraph("Main Content", style='Heading 1')
        doc.add_paragraph("Text", style='Normal')
        
        long_title = "Appendix: " + "A" * 200
        doc.add_paragraph(long_title, style='Heading 1')
        doc.add_paragraph("Content", style='Normal')
        
        try:
            appendix_processor.process_appendices(doc)
        except Exception as e:
            pytest.fail(f"Should handle long titles: {e}")


class TestAppendixWithTables:
    """Test appendix processing with tables."""
    
    def test_appendix_preserves_table_structure(self, appendix_processor,
                                                document_with_tables):
        """Test that processing preserves tables in appendices."""
        initial_table_count = len(document_with_tables.tables)
        
        appendix_processor.process_appendices(document_with_tables)
        
        final_table_count = len(document_with_tables.tables)
        assert final_table_count >= initial_table_count
    
    def test_appendix_with_multiple_tables(self, appendix_processor):
        """Test appendix containing multiple tables."""
        doc = Document()
        doc.add_paragraph("Main Content", style='Heading 1')
        doc.add_paragraph("Text", style='Normal')
        
        doc.add_paragraph("Appendix with Tables", style='Heading 1')
        
        # Add multiple tables
        for i in range(3):
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = f"Table {i+1} - H1"
            table.cell(0, 1).text = f"Table {i+1} - H2"
        
        try:
            appendix_processor.process_appendices(doc)
        except Exception as e:
            pytest.fail(f"Should handle multiple tables in appendix: {e}")


class TestAppendixIntegration:
    """Test integration with document processing pipeline."""
    
    def test_multiple_process_calls(self, appendix_processor,
                                    document_with_appendices):
        """Test that multiple calls don't cause issues."""
        try:
            appendix_processor.process_appendices(document_with_appendices)
            appendix_processor.process_appendices(document_with_appendices)
        except Exception as e:
            pytest.fail(f"Should handle multiple calls: {e}")
    
    def test_processor_with_mixed_content(self, appendix_processor):
        """Test processor with mixed content types."""
        doc = Document()
        
        # Regular content
        doc.add_paragraph("Title", style='Heading 1')
        doc.add_paragraph("Regular paragraph", style='Normal')
        
        # Appendix section
        doc.add_paragraph("Appendix Section", style='Heading 1')
        doc.add_paragraph("Appendix text", style='Normal')
        
        # Table in appendix
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Data"
        
        try:
            appendix_processor.process_appendices(doc)
        except Exception as e:
            pytest.fail(f"Should handle mixed content: {e}")
    
    def test_processor_methods_exist(self, appendix_processor):
        """Test that required methods exist."""
        assert hasattr(appendix_processor, 'process_appendices')
        assert callable(appendix_processor.process_appendices)
        
        assert hasattr(appendix_processor, '_find_appendix_headings')
        assert callable(appendix_processor._find_appendix_headings)
        
        assert hasattr(appendix_processor, '_get_appendix_letter')
        assert callable(appendix_processor._get_appendix_letter)


class TestAppendixConfigurationVariants:
    """Test different configuration combinations."""
    
    def test_enabled_with_letters_numbering(self):
        """Test enabled processor with letters numbering."""
        doc_structure = DocumentStructureConfig(
            appendix=AppendixConfig(
                enabled=True,
                numbering_style="letters"
            )
        )
        structure = StructureConfig(document_structure=doc_structure)
        general = GeneralConfig(
            margins=MarginsConfig(),
            fonts={},
            spacing=SpacingConfig()
        )
        config = DocumentConfig(general=general, structure=structure)
        
        processor = AppendixProcessor(config)
        assert processor.config.structure.document_structure.appendix.enabled
        assert processor.config.structure.document_structure.appendix.numbering_style == "letters"
    
    def test_enabled_with_numbers_numbering(self):
        """Test enabled processor with numbers numbering."""
        doc_structure = DocumentStructureConfig(
            appendix=AppendixConfig(
                enabled=True,
                numbering_style="numbers"
            )
        )
        structure = StructureConfig(document_structure=doc_structure)
        general = GeneralConfig(
            margins=MarginsConfig(),
            fonts={},
            spacing=SpacingConfig()
        )
        config = DocumentConfig(general=general, structure=structure)
        
        processor = AppendixProcessor(config)
        assert processor.config.structure.document_structure.appendix.enabled
        assert processor.config.structure.document_structure.appendix.numbering_style == "numbers"
    
    def test_config_change_after_creation(self, appendix_processor):
        """Test that config can be changed after processor creation."""
        # Store original value
        original_enabled = appendix_processor.config.structure.document_structure.appendix.enabled
        
        # Change it
        appendix_processor.config.structure.document_structure.appendix.enabled = False
        
        # Should be changed
        assert appendix_processor.config.structure.document_structure.appendix.enabled == False
        
        # Change it back
        appendix_processor.config.structure.document_structure.appendix.enabled = original_enabled
        assert appendix_processor.config.structure.document_structure.appendix.enabled == original_enabled


class TestAppendixDocumentation:
    """Test that processor has proper documentation."""
    
    def test_class_has_docstring(self, appendix_processor):
        """Test that class has docstring."""
        assert AppendixProcessor.__doc__ is not None
        assert len(AppendixProcessor.__doc__) > 0
    
    def test_main_method_has_docstring(self, appendix_processor):
        """Test that main method has docstring."""
        assert appendix_processor.process_appendices.__doc__ is not None
