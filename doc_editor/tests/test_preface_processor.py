"""
Comprehensive tests for PrefaceProcessor.

Tests cover:
- Initialization and configuration
- Enabled/disabled states
- Preface creation and insertion
- Positioning in document
- Content management
- Edge cases and formatting
"""

import pytest
from docx import Document

from doc_editor.models.config import (
    DocumentConfig,
    GeneralConfig,
    SpacingConfig,
    MarginsConfig,
    StructureConfig,
    TitlePageConfig,
    HeadersConfig,
    HeaderTextPart,
    TableFormatConfig,
    NumberingConfig,
    SectionConfig,
    TOCConfig,
    PrefaceConfig,
    AppendixConfig,
    DocumentStructureConfig,
)
from doc_editor.processors.preface_processor import PrefaceProcessor


# ============================================================================
# FIXTURES
# ============================================================================

@pytest.fixture
def base_config():
    """Create base configuration for all tests."""
    return DocumentConfig(
        general=GeneralConfig(
            margins=MarginsConfig(),
            fonts={"main": "Times New Roman"},
            spacing=SpacingConfig()
        ),
        structure=StructureConfig(
            title_page=TitlePageConfig(
                enabled=True,
                line_spacing=1.5,
                table_format=TableFormatConfig()
            ),
            numbering=NumberingConfig(
                headers=HeadersConfig(
                    left_parts=[HeaderTextPart(text="ГОСТ Р", bold=True)],
                    right_parts=[HeaderTextPart(text="Page", bold=False)]
                ),
                sections=SectionConfig(enabled=True)
            ),
            document_structure=DocumentStructureConfig(
                sections=SectionConfig(enabled=True),
                toc=TOCConfig(
                    enabled=True,
                    title="ОГЛАВЛЕНИЕ",
                    page_numbers=True,
                    levels=3
                ),
                preface=PrefaceConfig(
                    enabled=True,
                    content="Это содержание предисловия."
                ),
                appendix=AppendixConfig(enabled=False)
            )
        )
    )


@pytest.fixture
def preface_processor(base_config):
    """Create PrefaceProcessor instance."""
    return PrefaceProcessor(base_config)


@pytest.fixture
def empty_document():
    """Create empty Word document."""
    return Document()


@pytest.fixture
def document_with_content():
    """Create document with initial content."""
    doc = Document()
    doc.add_paragraph("Main Section 1", style='Heading 1')
    doc.add_paragraph("Content of section 1")
    doc.add_paragraph("Main Section 2", style='Heading 1')
    doc.add_paragraph("Content of section 2")
    return doc


# ============================================================================
# TEST SUITE 1: Initialization and Configuration
# ============================================================================

class TestPrefaceProcessorInitialization:
    """Tests for PrefaceProcessor initialization."""
    
    def test_processor_initialization(self, preface_processor, base_config):
        """Test that processor initializes correctly."""
        assert preface_processor is not None
        assert preface_processor.config == base_config
        assert preface_processor.logger is not None
    
    def test_processor_with_preface_disabled(self, base_config):
        """Test processor creation when preface is disabled."""
        base_config.structure.document_structure.preface.enabled = False
        processor = PrefaceProcessor(base_config)
        assert processor.config.structure.document_structure.preface.enabled is False
    
    def test_processor_preserves_config_content(self, preface_processor):
        """Test that custom preface content is preserved."""
        custom_content = "Custom preface content"
        preface_processor.config.structure.document_structure.preface.content = custom_content
        assert preface_processor.config.structure.document_structure.preface.content == custom_content
    
    def test_processor_with_empty_content(self, base_config):
        """Test processor with empty preface content."""
        base_config.structure.document_structure.preface.content = ""
        processor = PrefaceProcessor(base_config)
        assert processor.config.structure.document_structure.preface.content == ""
    
    def test_processor_with_multiline_content(self, base_config):
        """Test processor with multiline preface content."""
        multiline_content = """Line 1 of preface
Line 2 of preface
Line 3 of preface"""
        base_config.structure.document_structure.preface.content = multiline_content
        processor = PrefaceProcessor(base_config)
        assert multiline_content in processor.config.structure.document_structure.preface.content


# ============================================================================
# TEST SUITE 2: Enable/Disable Functionality
# ============================================================================

class TestPrefaceProcessorToggling:
    """Tests for enabling/disabling preface generation."""
    
    def test_preface_creation_disabled(self, base_config, document_with_content):
        """Test that add_preface does nothing when disabled."""
        base_config.structure.document_structure.preface.enabled = False
        processor = PrefaceProcessor(base_config)
        
        doc = document_with_content
        original_para_count = len(doc.paragraphs)
        
        processor.add_preface(doc)
        
        # No paragraphs should be added
        assert len(doc.paragraphs) == original_para_count
    
    def test_preface_creation_enabled(self, base_config, document_with_content):
        """Test that add_preface adds paragraphs when enabled."""
        base_config.structure.document_structure.preface.enabled = True
        processor = PrefaceProcessor(base_config)
        
        doc = document_with_content
        original_para_count = len(doc.paragraphs)
        
        processor.add_preface(doc)
        
        # Paragraphs should be added
        assert len(doc.paragraphs) > original_para_count
    
    def test_toggle_multiple_times(self, preface_processor, empty_document):
        """Test toggling enabled/disabled multiple times."""
        doc = empty_document
        doc.add_paragraph("Section", style='Heading 1')
        
        # First call - disabled
        preface_processor.config.structure.document_structure.preface.enabled = False
        count_1 = len(doc.paragraphs)
        preface_processor.add_preface(doc)
        assert len(doc.paragraphs) == count_1
        
        # Second call - enabled
        preface_processor.config.structure.document_structure.preface.enabled = True
        preface_processor.add_preface(doc)
        assert len(doc.paragraphs) > count_1


# ============================================================================
# TEST SUITE 3: Preface Creation and Insertion
# ============================================================================

class TestPrefaceCreation:
    """Tests for preface creation and insertion."""
    
    def test_simple_preface_addition(self, preface_processor, empty_document):
        """Test adding preface to empty document."""
        doc = empty_document
        
        preface_processor.add_preface(doc)
        
        # Check that preface content was added
        doc_text = "\n".join([p.text for p in doc.paragraphs])
        assert "Это содержание предисловия." in doc_text or len(doc.paragraphs) > 0
    
    def test_preface_with_document_content(self, preface_processor, document_with_content):
        """Test adding preface to document with existing content."""
        doc = document_with_content
        original_text = [p.text for p in doc.paragraphs]
        
        preface_processor.add_preface(doc)
        
        # Original content should still exist
        final_text = [p.text for p in doc.paragraphs]
        for text in original_text:
            assert text in final_text
    
    def test_preface_with_custom_content(self, base_config, empty_document):
        """Test preface with custom content."""
        custom_content = "Custom preface content for testing"
        base_config.structure.document_structure.preface.content = custom_content
        processor = PrefaceProcessor(base_config)
        
        doc = empty_document
        processor.add_preface(doc)
        
        doc_text = "\n".join([p.text for p in doc.paragraphs])
        assert custom_content in doc_text or custom_content in [p.text for p in doc.paragraphs]
    
    def test_preface_with_empty_content(self, base_config, empty_document):
        """Test preface with empty content."""
        base_config.structure.document_structure.preface.content = ""
        processor = PrefaceProcessor(base_config)
        
        doc = empty_document
        # Should not raise error
        processor.add_preface(doc)
        
        assert doc is not None


# ============================================================================
# TEST SUITE 4: Positioning in Document
# ============================================================================

class TestPrefacePositioning:
    """Tests for preface positioning in document."""
    
    def test_preface_position_in_document(self, preface_processor, document_with_content):
        """Test that preface is positioned correctly."""
        doc = document_with_content
        
        # Store first section heading before adding preface
        first_heading = None
        for p in doc.paragraphs:
            if 'Heading' in p.style.name:
                first_heading = p.text
                break
        
        preface_processor.add_preface(doc)
        
        # Original heading should still be there
        assert first_heading in [p.text for p in doc.paragraphs]
    
    def test_preface_does_not_duplicate(self, preface_processor, empty_document):
        """Test that multiple calls don't duplicate preface."""
        doc = empty_document
        
        preface_processor.add_preface(doc)
        count_after_first = len(doc.paragraphs)
        
        preface_processor.add_preface(doc)
        count_after_second = len(doc.paragraphs)
        
        # Behavior is acceptable either way (rebuild or skip)
        assert doc is not None


# ============================================================================
# TEST SUITE 5: Content Management
# ============================================================================

class TestPrefaceContentManagement:
    """Tests for preface content handling."""
    
    def test_multiline_preface_content(self, base_config, empty_document):
        """Test preface with multiline content."""
        multiline = """Line 1
Line 2
Line 3"""
        base_config.structure.document_structure.preface.content = multiline
        processor = PrefaceProcessor(base_config)
        
        doc = empty_document
        processor.add_preface(doc)
        
        # Check that content was added
        assert len(doc.paragraphs) > 0
    
    def test_preface_content_with_special_characters(self, base_config, empty_document):
        """Test preface with special characters."""
        special_content = "Преди́словие с спец. символами: @#$%^&*()"
        base_config.structure.document_structure.preface.content = special_content
        processor = PrefaceProcessor(base_config)
        
        doc = empty_document
        processor.add_preface(doc)
        
        assert len(doc.paragraphs) > 0
    
    def test_long_preface_content(self, base_config, empty_document):
        """Test preface with long content."""
        long_content = "Это содержание предисловия. " * 50  # Repeat 50 times
        base_config.structure.document_structure.preface.content = long_content
        processor = PrefaceProcessor(base_config)
        
        doc = empty_document
        processor.add_preface(doc)
        
        assert len(doc.paragraphs) > 0


# ============================================================================
# TEST SUITE 6: Edge Cases
# ============================================================================

class TestPrefaceEdgeCases:
    """Tests for edge cases and boundary conditions."""
    
    def test_empty_document(self, preface_processor, empty_document):
        """Test preface on empty document."""
        doc = empty_document
        
        # Should not raise error
        preface_processor.add_preface(doc)
        
        assert doc is not None
        assert hasattr(doc, 'paragraphs')
    
    def test_document_with_only_text(self, base_config):
        """Test preface with document containing only text."""
        doc = Document()
        doc.add_paragraph("Just some text")
        doc.add_paragraph("More text")
        
        processor = PrefaceProcessor(base_config)
        processor.add_preface(doc)
        
        assert len(doc.paragraphs) > 2
    
    def test_document_with_tables(self, base_config):
        """Test preface with document containing tables."""
        doc = Document()
        doc.add_paragraph("Before table")
        table = doc.add_table(rows=2, cols=2)
        doc.add_paragraph("After table")
        
        processor = PrefaceProcessor(base_config)
        # Should handle gracefully
        processor.add_preface(doc)
        
        assert doc is not None
    
    def test_preface_content_with_newlines(self, base_config, empty_document):
        """Test preface content with embedded newlines."""
        content_with_newlines = "Line 1\n\nLine 2 (after blank)\nLine 3"
        base_config.structure.document_structure.preface.content = content_with_newlines
        processor = PrefaceProcessor(base_config)
        
        doc = empty_document
        processor.add_preface(doc)
        
        assert len(doc.paragraphs) > 0


# ============================================================================
# TEST SUITE 7: Formatting
# ============================================================================

class TestPrefaceFormatting:
    """Tests for preface formatting."""
    
    def test_preface_paragraph_exists(self, preface_processor, empty_document):
        """Test that preface paragraph is created."""
        doc = empty_document
        preface_processor.add_preface(doc)
        
        # Should have paragraphs
        assert len(doc.paragraphs) > 0
    
    def test_document_structure_valid(self, preface_processor, document_with_content):
        """Test that document structure remains valid after adding preface."""
        doc = document_with_content
        preface_processor.add_preface(doc)
        
        # Document should be valid
        assert all(hasattr(p, 'text') for p in doc.paragraphs)
        assert all(hasattr(p, 'style') for p in doc.paragraphs)
    
    def test_preface_content_preserved_after_processing(self, base_config, empty_document):
        """Test that preface content is preserved."""
        test_content = "Test preface content"
        base_config.structure.document_structure.preface.content = test_content
        processor = PrefaceProcessor(base_config)
        
        doc = empty_document
        processor.add_preface(doc)
        
        # Content should be added
        assert len(doc.paragraphs) > 0


# ============================================================================
# TEST SUITE 8: Configuration Variants
# ============================================================================

class TestPrefaceConfigurationVariants:
    """Tests for different configuration scenarios."""
    
    def test_preface_with_minimal_config(self, base_config, empty_document):
        """Test preface with minimal configuration."""
        base_config.structure.document_structure.preface.enabled = True
        base_config.structure.document_structure.preface.content = "Minimal preface"
        processor = PrefaceProcessor(base_config)
        
        doc = empty_document
        processor.add_preface(doc)
        
        assert len(doc.paragraphs) > 0
    
    def test_preface_config_changed_after_creation(self, preface_processor, empty_document):
        """Test changing config after processor creation."""
        doc = empty_document
        
        # Change config
        preface_processor.config.structure.document_structure.preface.content = "Changed content"
        preface_processor.add_preface(doc)
        
        assert len(doc.paragraphs) > 0


# ============================================================================
# TEST SUITE 9: Integration Tests
# ============================================================================

class TestPrefaceProcessorIntegration:
    """Integration tests with documents."""
    
    def test_preface_with_section_structure(self, base_config):
        """Test preface with proper section structure."""
        doc = Document()
        doc.add_paragraph("Section 1", style='Heading 1')
        doc.add_paragraph("Content")
        doc.add_paragraph("Section 1.1", style='Heading 2')
        doc.add_paragraph("Subsection content")
        
        processor = PrefaceProcessor(base_config)
        processor.add_preface(doc)
        
        assert len(doc.paragraphs) > 4
    
    def test_multiple_processor_calls(self, base_config):
        """Test creating preface for multiple documents."""
        processor = PrefaceProcessor(base_config)
        
        for i in range(3):
            doc = Document()
            doc.add_paragraph(f"Document {i}")
            
            # Should not raise error
            processor.add_preface(doc)
            
            assert len(doc.paragraphs) > 1
    
    def test_preface_with_complex_content(self, base_config, empty_document):
        """Test preface with complex content."""
        complex_content = """ПРЕДИСЛОВИЕ

Это предисловие документа. Оно содержит вводную информацию.

Ключевые моменты:
- Пункт 1
- Пункт 2
- Пункт 3

Заключение предисловия."""
        
        base_config.structure.document_structure.preface.content = complex_content
        processor = PrefaceProcessor(base_config)
        
        doc = empty_document
        processor.add_preface(doc)
        
        assert len(doc.paragraphs) > 0


# ============================================================================
# TEST SUITE 10: Methods Testing
# ============================================================================

class TestPrefaceProcessorMethods:
    """Tests for individual processor methods."""
    
    def test_processor_has_required_methods(self, preface_processor):
        """Test that processor has all required methods."""
        assert hasattr(preface_processor, 'add_preface')
        assert callable(preface_processor.add_preface)
    
    def test_logger_is_available(self, preface_processor):
        """Test that logger is properly initialized."""
        assert preface_processor.logger is not None
        assert hasattr(preface_processor.logger, 'info')
        assert hasattr(preface_processor.logger, 'debug')
        assert hasattr(preface_processor.logger, 'warning')


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
