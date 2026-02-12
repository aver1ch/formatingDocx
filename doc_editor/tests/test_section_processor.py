"""
Тесты для SectionProcessor - многоуровневой нумерации разделов.
"""

import pytest
from docx import Document
from docx.oxml.ns import qn

from doc_editor.processors.section_processor import SectionProcessor
from doc_editor.models.config import (
    DocumentConfig,
    SectionConfig,
    DocumentStructureConfig,
    GeneralConfig,
    MarginsConfig,
    SpacingConfig,
    StructureConfig,
    NumberingConfig,
)


class TestSectionProcessor:
    """Тесты для SectionProcessor."""
    
    @pytest.fixture
    def section_config(self):
        """Создать конфигурацию с включенной нумерацией разделов."""
        return SectionConfig(
            enabled=True,
            start_number=1,
            numbering_format="decimal",
            include_in_toc=True,
            auto_number_headings=True,
            numbering_levels=3
        )
    
    @pytest.fixture
    def document_config(self, section_config):
        """Создать конфигурацию документа."""
        margins = MarginsConfig()
        spacing = SpacingConfig()
        general = GeneralConfig(
            margins=margins,
            fonts={},
            spacing=spacing
        )
        
        doc_structure = DocumentStructureConfig(sections=section_config)
        numbering = NumberingConfig()
        structure = StructureConfig(
            document_structure=doc_structure,
            numbering=numbering
        )
        
        config = DocumentConfig(general=general, structure=structure)
        return config
    
    @pytest.fixture
    def processor(self, document_config):
        """Создать процессор нумерации."""
        return SectionProcessor(document_config)
    
    @pytest.fixture
    def simple_document_with_headings(self):
        """Создать документ с простыми заголовками."""
        doc = Document()
        doc.add_paragraph("Введение", style='Heading 1')
        doc.add_paragraph("Общие положения", style='Heading 2')
        doc.add_paragraph("Область применения", style='Heading 2')
        doc.add_paragraph("Нормативные ссылки", style='Heading 3')
        return doc
    
    @pytest.fixture
    def complex_document_with_headings(self):
        """Создать сложный документ с вложенными заголовками."""
        doc = Document()
        doc.add_paragraph("Основные положения", style='Heading 1')
        doc.add_paragraph("Базовые концепции", style='Heading 2')
        doc.add_paragraph("Определение 1", style='Heading 3')
        doc.add_paragraph("Определение 2", style='Heading 3')
        doc.add_paragraph("Свойства", style='Heading 2')
        doc.add_paragraph("Методология", style='Heading 1')
        doc.add_paragraph("Подход", style='Heading 2')
        return doc
    
    def test_processor_initialization(self, processor):
        """Проверить инициализацию процессора."""
        assert processor.section_numbers == [0, 0, 0]
        assert processor.HEADING_LEVELS['Heading 1'] == 0
        assert processor.HEADING_LEVELS['Heading 2'] == 1
        assert processor.HEADING_LEVELS['Heading 3'] == 2
    
    def test_section_numbering_disabled(self, document_config):
        """Проверить, что нумерация не применяется, если отключена."""
        document_config.structure.document_structure.sections.enabled = False
        processor = SectionProcessor(document_config)
        
        doc = Document()
        doc.add_paragraph("Раздел 1", style='Heading 1')
        original_text = doc.paragraphs[0].text
        
        processor.apply_section_numbering(doc)
        
        # Текст не должен измениться
        assert doc.paragraphs[0].text == original_text
    
    def test_single_level_numbering(self, processor, simple_document_with_headings):
        """Проверить нумерацию первого уровня."""
        processor.apply_section_numbering(simple_document_with_headings)
        
        # Проверяем первый заголовок
        assert simple_document_with_headings.paragraphs[0].text.startswith("1 ")
    
    def test_two_level_numbering(self, processor, simple_document_with_headings):
        """Проверить нумерацию двух уровней."""
        processor.apply_section_numbering(simple_document_with_headings)
        
        paragraphs = simple_document_with_headings.paragraphs
        
        # Проверяем иерархию
        assert paragraphs[0].text.startswith("1 ")    # Heading 1
        assert paragraphs[1].text.startswith("1.1 ")  # Heading 2
        assert paragraphs[2].text.startswith("1.2 ")  # Heading 2
    
    def test_three_level_numbering(self, processor, simple_document_with_headings):
        """Проверить нумерацию трех уровней."""
        processor.apply_section_numbering(simple_document_with_headings)
        
        paragraphs = simple_document_with_headings.paragraphs
        
        # Проверяем третий уровень
        assert paragraphs[3].text.startswith("1.2.1 ")  # Heading 3
    
    def test_complex_document_structure(self, processor, complex_document_with_headings):
        """Проверить нумерацию сложного документа."""
        processor.apply_section_numbering(complex_document_with_headings)
        
        paragraphs = complex_document_with_headings.paragraphs
        
        # Проверяем корректность иерархии
        assert paragraphs[0].text.startswith("1 ")         # Heading 1
        assert paragraphs[1].text.startswith("1.1 ")       # Heading 2
        assert paragraphs[2].text.startswith("1.1.1 ")     # Heading 3
        assert paragraphs[3].text.startswith("1.1.2 ")     # Heading 3
        assert paragraphs[4].text.startswith("1.2 ")       # Heading 2 (сброс подуровней)
        assert paragraphs[5].text.startswith("2 ")         # Heading 1 (новый раздел)
        assert paragraphs[6].text.startswith("2.1 ")       # Heading 2 (сброс счетчиков)
    
    def test_section_numbers_reset(self, processor):
        """Проверить сброс счетчиков."""
        processor.section_numbers = [2, 3, 1]
        processor.reset_numbering()
        
        assert processor.section_numbers == [0, 0, 0]
    
    def test_get_section_number_level_0(self, processor):
        """Проверить формирование номера уровня 0."""
        processor.section_numbers = [5, 0, 0]
        assert processor._get_section_number(0) == "5"
    
    def test_get_section_number_level_1(self, processor):
        """Проверить формирование номера уровня 1."""
        processor.section_numbers = [5, 3, 0]
        assert processor._get_section_number(1) == "5.3"
    
    def test_get_section_number_level_2(self, processor):
        """Проверить формирование номера уровня 2."""
        processor.section_numbers = [5, 3, 2]
        assert processor._get_section_number(2) == "5.3.2"
    
    def test_already_numbered_heading(self, processor):
        """Проверить обработку уже нумерованного заголовка."""
        doc = Document()
        doc.add_paragraph("1 Раздел с номером", style='Heading 1')
        
        processor.apply_section_numbering(doc)
        
        # Должны иметь: "1 Раздел с номером" (парсится существующий номер)
        text = doc.paragraphs[0].text
        assert text.startswith("1 ")
    
    def test_mixed_heading_and_normal_text(self, processor):
        """Проверить обработку документа с обычным текстом и заголовками."""
        doc = Document()
        doc.add_paragraph("Раздел 1", style='Heading 1')
        doc.add_paragraph("Это обычный текст", style='Normal')
        doc.add_paragraph("Подраздел", style='Heading 2')
        
        processor.apply_section_numbering(doc)
        
        # Проверяем, что нумерация применена только к заголовкам
        assert doc.paragraphs[0].text.startswith("1 ")
        assert doc.paragraphs[1].text == "Это обычный текст"
        assert doc.paragraphs[2].text.startswith("1.1 ")
    
    def test_get_current_section_number_after_processing(self, processor, simple_document_with_headings):
        """Проверить получение текущего номера раздела после обработки."""
        processor.apply_section_numbering(simple_document_with_headings)
        
        # После обработки должен быть номер третьего раздела
        section_num = processor.get_current_section_number()
        assert section_num is not None
    
    def test_multiple_documents_in_sequence(self, processor):
        """Проверить обработку нескольких документов подряд."""
        # Первый документ
        doc1 = Document()
        doc1.add_paragraph("Раздел 1", style='Heading 1')
        processor.apply_section_numbering(doc1)
        assert doc1.paragraphs[0].text.startswith("1 ")
        
        # Сброс и второй документ
        processor.reset_numbering()
        doc2 = Document()
        doc2.add_paragraph("Раздел 1", style='Heading 1')
        processor.apply_section_numbering(doc2)
        assert doc2.paragraphs[0].text.startswith("1 ")
    
    def test_empty_document(self, processor):
        """Проверить обработку пустого документа."""
        doc = Document()
        
        # Не должно быть ошибок
        processor.apply_section_numbering(doc)
        assert len(doc.paragraphs) == 0
    
    def test_document_with_only_text(self, processor):
        """Проверить документ только с текстом (без заголовков)."""
        doc = Document()
        doc.add_paragraph("Просто текст 1", style='Normal')
        doc.add_paragraph("Просто текст 2", style='Normal')
        
        processor.apply_section_numbering(doc)
        
        # Текст не должен измениться
        assert doc.paragraphs[0].text == "Просто текст 1"
        assert doc.paragraphs[1].text == "Просто текст 2"
    
    def test_heading_formatting_preserved(self, processor, simple_document_with_headings):
        """Проверить, что форматирование заголовков сохраняется."""
        processor.apply_section_numbering(simple_document_with_headings)
        
        # Первый параграф - это Heading 1
        first_para = simple_document_with_headings.paragraphs[0]
        
        # Проверяем, что стиль сохранился
        assert first_para.style.name == 'Heading 1'
        
        # Проверяем, что текст нумерован
        assert first_para.text.startswith("1 ")


class TestSectionProcessorIntegration:
    """Интеграционные тесты для SectionProcessor."""
    
    def test_section_processor_with_real_document_structure(self):
        """Проверить работу с реальной структурой документа."""
        # Создаем конфигурацию
        section_config = SectionConfig(enabled=True)
        doc_structure = DocumentStructureConfig(sections=section_config)
        numbering = NumberingConfig()
        
        margins = MarginsConfig()
        spacing = SpacingConfig()
        general = GeneralConfig(
            margins=margins,
            fonts={'main': {'family': 'Arial'}},
            spacing=spacing
        )
        structure = StructureConfig(
            document_structure=doc_structure,
            numbering=numbering
        )
        config = DocumentConfig(general=general, structure=structure)
        
        # Создаем документ
        doc = Document()
        doc.add_paragraph("Введение", style='Heading 1')
        doc.add_paragraph("Некоторый текст")
        doc.add_paragraph("Методология", style='Heading 1')
        doc.add_paragraph("Подход", style='Heading 2')
        doc.add_paragraph("Детали", style='Heading 3')
        
        # Применяем обработку
        processor = SectionProcessor(config)
        processor.apply_section_numbering(doc)
        
        # Проверяем результаты
        headings = [p for p in doc.paragraphs if p.style.name.startswith('Heading')]
        assert len(headings) == 4
        assert headings[0].text.startswith("1 ")
        assert headings[1].text.startswith("2 ")
        assert headings[2].text.startswith("2.1 ")
        assert headings[3].text.startswith("2.1.1 ")
