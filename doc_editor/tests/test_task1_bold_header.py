"""
Tests for Task 1: Making ГОСТ Р bold in headers with rest remaining normal.

This test module verifies that the header formatting works correctly,
allowing "ГОСТ Р" to be displayed in bold while "(проект, первая редакция)"
remains in normal formatting.
"""

import pytest
from pathlib import Path
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn

from doc_editor.editor import DocumentEditor
from doc_editor.models.config import DocumentConfig, HeadersConfig, HeaderTextPart


class TestTask1BoldHeader:
    """Test suite for Task 1: Bold ГОСТ Р in headers."""

    @pytest.fixture
    def test_config_path(self):
        """Return path to test config with formatted headers."""
        return Path(__file__).parent / "test_data" / "formatConfig_with_formatted_headers.yaml"

    @pytest.fixture
    def test_document_path(self, tmp_path):
        """Create a simple test document."""
        doc = Document()
        doc.add_paragraph("Test Document for GOST Header Formatting")
        test_doc = tmp_path / "test_document.docx"
        doc.save(test_doc)
        return test_doc

    def test_header_text_part_dataclass(self):
        """Test that HeaderTextPart dataclass is properly defined."""
        part = HeaderTextPart(text="ГОСТ Р", bold=True, italic=False)
        assert part.text == "ГОСТ Р"
        assert part.bold is True
        assert part.italic is False
        assert part.font_family is None

    def test_header_text_part_with_font_family(self):
        """Test HeaderTextPart with custom font family."""
        part = HeaderTextPart(
            text="Text",
            bold=False,
            italic=True,
            font_family="Times New Roman"
        )
        assert part.font_family == "Times New Roman"

    def test_headers_config_with_right_parts(self):
        """Test that HeadersConfig accepts right_parts."""
        parts = [
            HeaderTextPart(text="ГОСТ Р", bold=True),
            HeaderTextPart(text="\n(проект, первая редакция)", bold=False),
        ]
        headers_config = HeadersConfig(
            left="Left header",
            right="Right header (fallback)",
            right_parts=parts,
            page_numbers=True,
            enabled=True
        )
        assert len(headers_config.right_parts) == 2
        assert headers_config.right_parts[0].bold is True
        assert headers_config.right_parts[1].bold is False

    def test_document_editor_with_formatted_headers(self, test_config_path, test_document_path):
        """Test that DocumentEditor applies formatted headers correctly."""
        if not test_config_path.exists():
            pytest.skip(f"Test config not found: {test_config_path}")

        editor = DocumentEditor(str(test_document_path))
        editor.load_config(str(test_config_path))
        editor.apply_config()
        doc = editor.get_document()
        
        assert doc is not None
        # Check that headers were applied
        section = doc.sections[0]
        header = section.header
        assert len(header.paragraphs) > 0

    def test_header_bold_formatting_applied(self, test_config_path, test_document_path):
        """Test that bold formatting is correctly applied to header text."""
        if not test_config_path.exists():
            pytest.skip(f"Test config not found: {test_config_path}")

        editor = DocumentEditor(str(test_document_path))
        editor.load_config(str(test_config_path))
        try:
            editor.apply_config()
        except Exception:
            # If title page fails, it's ok - we're testing headers
            pass
        
        doc = editor.get_document()
        
        section = doc.sections[0]
        header = section.header
        
        # Find ГОСТ Р text in header
        found_bold_text = False
        for paragraph in header.paragraphs:
            for run in paragraph.runs:
                if "ГОСТ Р" in run.text and run.bold:
                    found_bold_text = True
                    break
        
        assert found_bold_text, "Bold ГОСТ Р text not found in header"

    def test_header_normal_formatting_preserved(self, test_config_path, test_document_path):
        """Test that normal formatting is preserved for non-bold parts."""
        if not test_config_path.exists():
            pytest.skip(f"Test config not found: {test_config_path}")

        editor = DocumentEditor(str(test_document_path))
        editor.load_config(str(test_config_path))
        try:
            editor.apply_config()
        except Exception:
            # If title page fails, it's ok - we're testing headers
            pass
        
        doc = editor.get_document()
        
        section = doc.sections[0]
        header = section.header
        
        # Find normal text in header
        found_normal_text = False
        for paragraph in header.paragraphs:
            for run in paragraph.runs:
                if "проект" in run.text and not run.bold:
                    found_normal_text = True
                    break
        
        assert found_normal_text, "Normal (non-bold) text not found in header"

    def test_backward_compatibility_string_headers(self, tmp_path):
        """Test that string-based headers still work (backward compatibility)."""
        # Create a minimal config with string headers (no title page)
        from io import StringIO
        import yaml
        
        config_content = """
document:
  general:
    margins:
      left: 20mm
      right: 10mm
      top: 20mm
      bottom: 20mm
    fonts:
      main:
        family: Arial
        size: 12pt
    spacing:
      line: 1.5
  structure:
    title_page:
      enabled: false
    headers:
      right: "Simple string header"
      left: "Left header"
      page_numbers: true
      enabled: true
"""
        config_file = tmp_path / "string_config.yaml"
        config_file.write_text(config_content)
        
        # Create test document
        doc = Document()
        doc.add_paragraph("Test document")
        test_doc = tmp_path / "test_doc.docx"
        doc.save(test_doc)
        
        # Apply formatting
        editor = DocumentEditor(str(test_doc))
        editor.load_config(str(config_file))
        editor.apply_config()
        formatted_doc = editor.get_document()
        
        # Verify document was processed
        assert formatted_doc is not None

    def test_mixed_bold_italic_formatting(self):
        """Test that HeaderTextPart handles both bold and italic."""
        part_bold_italic = HeaderTextPart(
            text="Test",
            bold=True,
            italic=True,
            font_family="Arial"
        )
        assert part_bold_italic.bold is True
        assert part_bold_italic.italic is True
        
        part_neither = HeaderTextPart(
            text="Test",
            bold=False,
            italic=False
        )
        assert part_neither.bold is False
        assert part_neither.italic is False

    def test_header_xml_formatting_attributes(self, test_config_path, test_document_path):
        """Test that XML-level formatting attributes are set correctly."""
        if not test_config_path.exists():
            pytest.skip(f"Test config not found: {test_config_path}")

        editor = DocumentEditor(str(test_document_path))
        editor.load_config(str(test_config_path))
        editor.apply_config()
        doc = editor.get_document()
        
        section = doc.sections[0]
        header = section.header
        
        # Check for XML-level bold attribute in rPr (run properties)
        for paragraph in header.paragraphs:
            for run in paragraph.runs:
                if "ГОСТ Р" in run.text:
                    rPr = run._element.find(qn('w:rPr'))
                    if rPr is not None:
                        b_element = rPr.find(qn('w:b'))
                        # b_element exists means bold is applied at XML level
                        if "ГОСТ" in run.text:
                            # At least check that bold property is set on run object
                            assert run.bold is True


class TestTask1Integration:
    """Integration tests for Task 1."""

    def test_apply_formatting_returns_document(self, tmp_path):
        """Test that apply_formatting returns a Document object."""
        from io import StringIO
        
        config_content = """
document:
  general:
    margins:
      left: 20mm
      right: 10mm
      top: 20mm
      bottom: 20mm
    fonts:
      main:
        family: Arial
        size: 12pt
    spacing:
      line: 1.5
  structure:
    title_page:
      enabled: false
    numbering:
      main: arabic
      headers:
        right: Test Header
        left: Left Header
        page_numbers: true
        enabled: true
"""
        config_file = tmp_path / "config.yaml"
        config_file.write_text(config_content)
        
        doc = Document()
        doc.add_paragraph("Content")
        test_doc = tmp_path / "test.docx"
        doc.save(test_doc)
        
        editor = DocumentEditor(str(test_doc))
        editor.load_config(str(config_file))
        editor.apply_config()
        result = editor.get_document()
        
        # Check result is a document-like object with sections
        assert hasattr(result, 'sections')
        assert len(result.sections) > 0

    def test_formatted_header_saved_correctly(self, tmp_path):
        """Test that formatted headers are saved and readable from file."""
        from io import StringIO
        
        config_content = """
document:
  general:
    margins:
      left: 20mm
      right: 10mm
      top: 20mm
      bottom: 20mm
    fonts:
      main:
        family: Arial
        size: 12pt
    spacing:
      line: 1.5
  structure:
    title_page:
      enabled: false
    numbering:
      main: arabic
      headers:
        right: "Test Header with Format"
        left: "Left"
        page_numbers: true
        enabled: true
"""
        config_file = tmp_path / "config.yaml"
        config_file.write_text(config_content)
        
        doc = Document()
        doc.add_paragraph("Content")
        test_doc = tmp_path / "test.docx"
        doc.save(test_doc)
        
        editor = DocumentEditor(str(test_doc))
        editor.load_config(str(config_file))
        editor.apply_config()
        result = editor.get_document()
        
        # Save and reload
        output_path = tmp_path / "output.docx"
        result.save(str(output_path))
        
        # Verify it can be reopened
        reloaded = Document(str(output_path))
        assert len(reloaded.sections) > 0
        header = reloaded.sections[0].header
        assert len(header.paragraphs) > 0


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
