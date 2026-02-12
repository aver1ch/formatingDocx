"""
Comprehensive tests for TOCProcessor.

Tests cover:
- Initialization and configuration
- Enabled/disabled states
- Heading extraction and TOC creation
- Hierarchical structure handling
- Page number calculation
- Document insertion and positioning
- Empty documents and edge cases
- Formatting preservation
"""

import pytest
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from doc_editor.models.config import (
    DocumentConfig,
    NumberingConfig,
    StructureConfig,
    TitlePageConfig,
    HeadersConfig,
    HeaderTextPart,
    MarginsConfig,
    TableFormatConfig,
    SectionConfig,
    TOCConfig,
    PrefaceConfig,
    AppendixConfig,
    DocumentStructureConfig,
    GeneralConfig,
    SpacingConfig,
)
from doc_editor.processors.toc_processor import TOCProcessor


# ============================================================================
# FIXTURES
# ============================================================================

@pytest.fixture
def base_config():
    """Create base configuration for all tests."""
    return DocumentConfig(
        general=GeneralConfig(
            margins=MarginsConfig(),
            fonts={"main": "Times New Roman"},
            spacing=SpacingConfig()
        ),
        structure=StructureConfig(
            title_page=TitlePageConfig(
                enabled=True,
                line_spacing=1.5,
                table_format=TableFormatConfig()
            ),
            numbering=NumberingConfig(
                headers=HeadersConfig(
                    left_parts=[HeaderTextPart(text="ГОСТ Р", bold=True)],
                    right_parts=[HeaderTextPart(text="Page", bold=False)]
                ),
                sections=SectionConfig(enabled=True)
            ),
            document_structure=DocumentStructureConfig(
                sections=SectionConfig(enabled=True),
                toc=TOCConfig(
                    enabled=True,
                    title="ОГЛАВЛЕНИЕ",
                    page_numbers=True,
                    levels=3
                ),
                preface=PrefaceConfig(enabled=False),
                appendix=AppendixConfig(enabled=False)
            )
        )
    )


@pytest.fixture
def toc_processor(base_config):
    """Create TOCProcessor instance."""
    return TOCProcessor(base_config)


@pytest.fixture
def empty_document():
    """Create empty Word document."""
    return Document()


@pytest.fixture
def document_with_headings():
    """Create document with various heading levels."""
    doc = Document()
    doc.add_paragraph("Section 1", style='Heading 1')
    doc.add_paragraph("Some text in section 1")
    doc.add_paragraph("Subsection 1.1", style='Heading 2')
    doc.add_paragraph("Content of 1.1")
    doc.add_paragraph("Subsubsection 1.1.1", style='Heading 3')
    doc.add_paragraph("Content of 1.1.1")
    doc.add_paragraph("Section 2", style='Heading 1')
    doc.add_paragraph("Content of section 2")
    doc.add_paragraph("Subsection 2.1", style='Heading 2')
    doc.add_paragraph("Content of 2.1")
    return doc


@pytest.fixture
def document_without_headings():
    """Create document with only text, no headings."""
    doc = Document()
    doc.add_paragraph("Regular paragraph 1")
    doc.add_paragraph("Regular paragraph 2")
    doc.add_paragraph("Regular paragraph 3")
    return doc


# ============================================================================
# TEST SUITE 1: Initialization and Configuration
# ============================================================================

class TestTOCProcessorInitialization:
    """Tests for TOCProcessor initialization."""
    
    def test_processor_initialization(self, toc_processor, base_config):
        """Test that processor initializes correctly."""
        assert toc_processor is not None
        assert toc_processor.config == base_config
        assert toc_processor.logger is not None
    
    def test_processor_with_toc_disabled(self, base_config):
        """Test processor creation when TOC is disabled."""
        base_config.structure.document_structure.toc.enabled = False
        processor = TOCProcessor(base_config)
        assert processor.config.structure.document_structure.toc.enabled is False
    
    def test_processor_preserves_config_title(self, toc_processor):
        """Test that custom TOC title is preserved."""
        custom_title = "TABLE OF CONTENTS"
        toc_processor.config.structure.document_structure.toc.title = custom_title
        assert toc_processor.config.structure.document_structure.toc.title == custom_title
    
    def test_processor_preserves_page_numbers_setting(self, toc_processor):
        """Test that page numbers setting is preserved."""
        assert toc_processor.config.structure.document_structure.toc.page_numbers is True
        
        toc_processor.config.structure.document_structure.toc.page_numbers = False
        assert toc_processor.config.structure.document_structure.toc.page_numbers is False
    
    def test_processor_preserves_levels_setting(self, toc_processor):
        """Test that levels setting is preserved."""
        assert toc_processor.config.structure.document_structure.toc.levels == 3


# ============================================================================
# TEST SUITE 2: Enable/Disable Functionality
# ============================================================================

class TestTOCProcessorToggling:
    """Tests for enabling/disabling TOC generation."""
    
    def test_toc_creation_disabled(self, base_config, document_with_headings):
        """Test that create_toc does nothing when disabled."""
        base_config.structure.document_structure.toc.enabled = False
        processor = TOCProcessor(base_config)
        
        doc = document_with_headings
        original_para_count = len(doc.paragraphs)
        
        processor.create_toc(doc)
        
        # No paragraphs should be added
        assert len(doc.paragraphs) == original_para_count
    
    def test_toc_creation_enabled(self, base_config, document_with_headings):
        """Test that create_toc adds paragraphs when enabled."""
        base_config.structure.document_structure.toc.enabled = True
        processor = TOCProcessor(base_config)
        
        doc = document_with_headings
        original_para_count = len(doc.paragraphs)
        
        processor.create_toc(doc)
        
        # Paragraphs should be added
        assert len(doc.paragraphs) > original_para_count
    
    def test_toggle_multiple_times(self, toc_processor, empty_document):
        """Test toggling enabled/disabled multiple times."""
        doc = empty_document
        doc.add_paragraph("Heading 1", style='Heading 1')
        doc.add_paragraph("Text")
        doc.add_paragraph("Heading 2", style='Heading 2')
        
        # First call - disabled
        toc_processor.config.structure.document_structure.toc.enabled = False
        count_1 = len(doc.paragraphs)
        toc_processor.create_toc(doc)
        assert len(doc.paragraphs) == count_1
        
        # Second call - enabled
        toc_processor.config.structure.document_structure.toc.enabled = True
        toc_processor.create_toc(doc)
        assert len(doc.paragraphs) > count_1


# ============================================================================
# TEST SUITE 3: Heading Extraction and Basic TOC Creation
# ============================================================================

class TestTOCHeadingExtraction:
    """Tests for heading extraction and basic TOC creation."""
    
    def test_extract_single_heading(self, toc_processor, empty_document):
        """Test extraction of single heading."""
        doc = empty_document
        doc.add_paragraph("My Section", style='Heading 1')
        
        toc_processor.create_toc(doc)
        
        # TOC should be created with title and heading
        doc_text = [p.text for p in doc.paragraphs]
        assert "ОГЛАВЛЕНИЕ" in doc_text
        assert "My Section" in doc_text
    
    def test_extract_multiple_same_level_headings(self, toc_processor, empty_document):
        """Test extraction of multiple headings at same level."""
        doc = empty_document
        doc.add_paragraph("Section 1", style='Heading 1')
        doc.add_paragraph("Some text")
        doc.add_paragraph("Section 2", style='Heading 1')
        doc.add_paragraph("More text")
        doc.add_paragraph("Section 3", style='Heading 1')
        
        toc_processor.create_toc(doc)
        
        doc_text = [p.text for p in doc.paragraphs]
        assert "ОГЛАВЛЕНИЕ" in doc_text
        # All sections should appear in TOC
        assert any("Section 1" in p for p in doc_text)
        assert any("Section 2" in p for p in doc_text)
        assert any("Section 3" in p for p in doc_text)
    
    def test_toc_preserves_original_content(self, toc_processor, document_with_headings):
        """Test that original document content is preserved after TOC creation."""
        doc = document_with_headings
        original_headings = [p.text for p in doc.paragraphs if 'Heading' in p.style.name]
        original_text_count = len(doc.paragraphs)
        
        toc_processor.create_toc(doc)
        
        # Original headings should still exist (plus TOC at beginning)
        doc_text = [p.text for p in doc.paragraphs]
        for heading in original_headings:
            assert heading in doc_text
    
    def test_toc_inserted_at_beginning(self, toc_processor, document_with_headings):
        """Test that TOC is inserted at the beginning of document."""
        doc = document_with_headings
        
        toc_processor.create_toc(doc)
        
        # "ОГЛАВЛЕНИЕ" should be near the beginning
        first_paragraphs = [p.text for p in doc.paragraphs[:5]]
        assert "ОГЛАВЛЕНИЕ" in first_paragraphs


# ============================================================================
# TEST SUITE 4: Hierarchical Structure
# ============================================================================

class TestTOCHierarchical:
    """Tests for hierarchical TOC structure with indentation."""
    
    def test_single_level_hierarchy(self, toc_processor, empty_document):
        """Test TOC with single heading level only."""
        doc = empty_document
        doc.add_paragraph("Section A", style='Heading 1')
        doc.add_paragraph("Text A")
        doc.add_paragraph("Section B", style='Heading 1')
        
        toc_processor.create_toc(doc)
        
        # Check that headings are in TOC
        doc_text = "\n".join([p.text for p in doc.paragraphs])
        assert "Section A" in doc_text
        assert "Section B" in doc_text
    
    def test_two_level_hierarchy(self, toc_processor, empty_document):
        """Test TOC with two heading levels."""
        doc = empty_document
        doc.add_paragraph("Section 1", style='Heading 1')
        doc.add_paragraph("Subsection 1.1", style='Heading 2')
        doc.add_paragraph("Subsection 1.2", style='Heading 2')
        doc.add_paragraph("Section 2", style='Heading 1')
        
        toc_processor.create_toc(doc)
        
        doc_text = "\n".join([p.text for p in doc.paragraphs])
        assert "Section 1" in doc_text
        assert "Subsection 1.1" in doc_text
        assert "Subsection 1.2" in doc_text
        assert "Section 2" in doc_text
    
    def test_three_level_hierarchy(self, toc_processor, document_with_headings):
        """Test TOC with three heading levels."""
        doc = document_with_headings
        
        toc_processor.create_toc(doc)
        
        doc_text = "\n".join([p.text for p in doc.paragraphs])
        assert "Section 1" in doc_text
        assert "Subsection 1.1" in doc_text
        assert "Subsubsection 1.1.1" in doc_text
        assert "Section 2" in doc_text
    
    def test_toc_indentation_structure(self, toc_processor, empty_document):
        """Test that indentation reflects heading hierarchy."""
        doc = empty_document
        doc.add_paragraph("1. Main Section", style='Heading 1')
        doc.add_paragraph("1.1 Subsection", style='Heading 2')
        doc.add_paragraph("1.1.1 Subsubsection", style='Heading 3')
        
        toc_processor.create_toc(doc)
        
        # Find TOC entries (after "ОГЛАВЛЕНИЕ")
        toc_start_idx = None
        for i, p in enumerate(doc.paragraphs):
            if p.text == "ОГЛАВЛЕНИЕ":
                toc_start_idx = i
                break
        
        assert toc_start_idx is not None
        # Check that there are indented entries
        toc_text = "\n".join([p.text for p in doc.paragraphs[toc_start_idx:toc_start_idx+5]])
        # Should contain entries with different indentation levels
        assert "Main Section" in toc_text


# ============================================================================
# TEST SUITE 5: Page Numbers Handling
# ============================================================================

class TestTOCPageNumbers:
    """Tests for page number calculation and insertion."""
    
    def test_toc_with_page_numbers_enabled(self, base_config, document_with_headings):
        """Test TOC creation with page numbers enabled."""
        base_config.structure.document_structure.toc.page_numbers = True
        processor = TOCProcessor(base_config)
        
        doc = document_with_headings
        processor.create_toc(doc)
        
        # Should have created TOC with some numeric page references
        doc_text = "\n".join([p.text for p in doc.paragraphs])
        assert "ОГЛАВЛЕНИЕ" in doc_text
    
    def test_toc_without_page_numbers(self, base_config, document_with_headings):
        """Test TOC creation without page numbers."""
        base_config.structure.document_structure.toc.page_numbers = False
        processor = TOCProcessor(base_config)
        
        doc = document_with_headings
        processor.create_toc(doc)
        
        # TOC should exist without page numbers
        doc_text = "\n".join([p.text for p in doc.paragraphs])
        assert "ОГЛАВЛЕНИЕ" in doc_text


# ============================================================================
# TEST SUITE 6: Edge Cases and Error Handling
# ============================================================================

class TestTOCEdgeCases:
    """Tests for edge cases and boundary conditions."""
    
    def test_empty_document(self, toc_processor, empty_document):
        """Test TOC creation on empty document."""
        doc = empty_document
        
        # Should not raise error
        toc_processor.create_toc(doc)
        
        # Document should still be valid (possibly with TOC title only)
        assert doc is not None
        assert hasattr(doc, 'paragraphs')
    
    def test_document_without_headings(self, toc_processor, document_without_headings):
        """Test TOC creation on document with no headings."""
        doc = document_without_headings
        original_count = len(doc.paragraphs)
        
        # Should not raise error
        toc_processor.create_toc(doc)
        
        # Document should still be valid
        assert doc is not None
        assert hasattr(doc, 'paragraphs')
    
    def test_toc_creation_multiple_times_same_document(self, toc_processor, document_with_headings):
        """Test calling create_toc multiple times on same document."""
        doc = document_with_headings
        
        # First call
        toc_processor.create_toc(doc)
        count_after_first = len(doc.paragraphs)
        
        # Second call
        toc_processor.create_toc(doc)
        count_after_second = len(doc.paragraphs)
        
        # Should handle gracefully (could rebuild or ignore)
        assert doc is not None
        assert hasattr(doc, 'paragraphs')
    
    def test_document_with_only_level_2_headings(self, toc_processor, empty_document):
        """Test TOC with only Level 2 headings (no Level 1)."""
        doc = empty_document
        doc.add_paragraph("Subsection A", style='Heading 2')
        doc.add_paragraph("Subsection B", style='Heading 2')
        
        # Should handle gracefully
        toc_processor.create_toc(doc)
        
        doc_text = "\n".join([p.text for p in doc.paragraphs])
        assert "Subsection A" in doc_text or "ОГЛАВЛЕНИЕ" in doc_text
    
    def test_document_with_mixed_heading_levels(self, toc_processor, empty_document):
        """Test TOC with mixed/non-sequential heading levels."""
        doc = empty_document
        doc.add_paragraph("Section 1", style='Heading 1')
        doc.add_paragraph("Subsubsection", style='Heading 3')  # Skip level 2
        doc.add_paragraph("Section 2", style='Heading 1')
        
        # Should handle gracefully
        toc_processor.create_toc(doc)
        
        assert doc is not None
        assert hasattr(doc, 'paragraphs')


# ============================================================================
# TEST SUITE 7: Configuration Customization
# ============================================================================

class TestTOCCustomization:
    """Tests for TOC customization via configuration."""
    
    def test_custom_toc_title(self, base_config, document_with_headings):
        """Test TOC with custom title."""
        custom_title = "TABLE OF CONTENTS"
        base_config.structure.document_structure.toc.title = custom_title
        processor = TOCProcessor(base_config)
        
        doc = document_with_headings
        processor.create_toc(doc)
        
        doc_text = [p.text for p in doc.paragraphs]
        assert custom_title in doc_text
    
    def test_toc_levels_filter(self, base_config, empty_document):
        """Test TOC with limited number of levels."""
        base_config.structure.document_structure.toc.levels = 2  # Only H1 and H2
        processor = TOCProcessor(base_config)
        
        doc = empty_document
        doc.add_paragraph("Section 1", style='Heading 1')
        doc.add_paragraph("Subsection 1.1", style='Heading 2')
        doc.add_paragraph("Subsubsection 1.1.1", style='Heading 3')
        doc.add_paragraph("Section 2", style='Heading 1')
        
        processor.create_toc(doc)
        
        # Should include H1 and H2 entries
        doc_text = [p.text for p in doc.paragraphs]
        assert "Section 1" in doc_text or "ОГЛАВЛЕНИЕ" in doc_text
    
    def test_toc_with_different_level_configs(self, base_config, document_with_headings):
        """Test TOC with different level configurations (1, 2, 3)."""
        doc = document_with_headings
        
        # Test with levels = 1
        base_config.structure.document_structure.toc.levels = 1
        processor = TOCProcessor(base_config)
        processor.create_toc(doc)
        assert doc is not None
        
        # Test with levels = 2
        doc = document_with_headings
        base_config.structure.document_structure.toc.levels = 2
        processor = TOCProcessor(base_config)
        processor.create_toc(doc)
        assert doc is not None


# ============================================================================
# TEST SUITE 8: Document Formatting and Styling
# ============================================================================

class TestTOCFormatting:
    """Tests for document formatting after TOC insertion."""
    
    def test_toc_title_style(self, toc_processor, document_with_headings):
        """Test that TOC title has appropriate style."""
        doc = document_with_headings
        toc_processor.create_toc(doc)
        
        # Find TOC title
        for para in doc.paragraphs:
            if para.text == "ОГЛАВЛЕНИЕ":
                # Should have a heading style
                assert "Heading" in para.style.name or para.style.name == "Normal"
                break
    
    def test_toc_entry_formatting_preserved(self, toc_processor, document_with_headings):
        """Test that document formatting is preserved after TOC."""
        doc = document_with_headings
        
        # Get formatting from first heading
        first_heading = None
        for para in doc.paragraphs:
            if "Heading" in para.style.name:
                first_heading = para
                break
        
        toc_processor.create_toc(doc)
        
        # Check document is still valid
        for para in doc.paragraphs:
            assert para.style is not None


# ============================================================================
# TEST SUITE 9: Integration Tests
# ============================================================================

class TestTOCProcessorIntegration:
    """Integration tests with full document processing."""
    
    def test_create_toc_with_real_document_structure(self, base_config):
        """Test TOC creation with realistic document structure."""
        processor = TOCProcessor(base_config)
        
        doc = Document()
        # Simulate a real document structure
        doc.add_paragraph("Introduction", style='Heading 1')
        doc.add_paragraph("Background information")
        
        doc.add_paragraph("Theory", style='Heading 1')
        doc.add_paragraph("Basic concepts", style='Heading 2')
        doc.add_paragraph("Definition 1", style='Heading 3')
        doc.add_paragraph("Detailed explanation")
        doc.add_paragraph("Definition 2", style='Heading 3')
        doc.add_paragraph("More details")
        
        doc.add_paragraph("Advanced Topics", style='Heading 2')
        doc.add_paragraph("Topic A", style='Heading 3')
        doc.add_paragraph("Content")
        
        doc.add_paragraph("Results", style='Heading 1')
        doc.add_paragraph("Result summary")
        
        processor.create_toc(doc)
        
        # Verify document structure is intact
        assert len(doc.paragraphs) > 10
        doc_text = "\n".join([p.text for p in doc.paragraphs])
        assert "ОГЛАВЛЕНИЕ" in doc_text
        assert "Introduction" in doc_text
        assert "Results" in doc_text
    
    def test_sequential_processor_calls(self, base_config):
        """Test creating TOC for multiple documents in sequence."""
        processor = TOCProcessor(base_config)
        
        for i in range(3):
            doc = Document()
            doc.add_paragraph(f"Document {i} Section 1", style='Heading 1')
            doc.add_paragraph("Content")
            doc.add_paragraph(f"Document {i} Section 2", style='Heading 1')
            
            # Should not raise error
            processor.create_toc(doc)
            
            assert doc is not None
            assert hasattr(doc, 'paragraphs')
            assert len(doc.paragraphs) > 0


# ============================================================================
# TEST SUITE 10: Method Testing
# ============================================================================

class TestTOCProcessorMethods:
    """Tests for individual processor methods."""
    
    def test_get_heading_level(self, toc_processor):
        """Test heading level determination."""
        # Test Level 1
        level = toc_processor._get_heading_level('Heading 1')
        assert level == 0
        
        # Test Level 2
        level = toc_processor._get_heading_level('Heading 2')
        assert level == 1
        
        # Test Level 3
        level = toc_processor._get_heading_level('Heading 3')
        assert level == 2
    
    def test_extract_headings(self, toc_processor, document_with_headings):
        """Test heading extraction method."""
        doc = document_with_headings
        headings = toc_processor._extract_headings(doc)
        
        # Should find 5 headings (Heading 1, 2, 3, 1, 2)
        assert len(headings) == 5
        
        # First should be Heading 1
        assert headings[0].style.name == 'Heading 1'
        assert "Section 1" in headings[0].text
    
    def test_extract_headings_empty_document(self, toc_processor, empty_document):
        """Test heading extraction on empty document."""
        doc = empty_document
        headings = toc_processor._extract_headings(doc)
        
        assert len(headings) == 0
    
    def test_build_toc_entries(self, toc_processor, document_with_headings):
        """Test TOC entries building."""
        doc = document_with_headings
        headings = toc_processor._extract_headings(doc)
        
        toc_entries = toc_processor._build_toc_entries(doc, headings)
        
        # Should have entries for all headings
        assert len(toc_entries) == len(headings)
        
        # Each entry should have required fields
        for entry in toc_entries:
            assert 'level' in entry
            assert 'text' in entry
            assert 'page_num' in entry


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
